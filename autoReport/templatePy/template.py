import os
import sys
import re
import datetime
import pandas as pd
# import modin.pandas as pd
from collections import Counter
# from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
# import win32com.client


class Template:	
	try:
		projectinfo = pd.read_excel('project_information.xls', header=None)
		species = projectinfo.iloc[3, 1]
		groupvs = projectinfo.iloc[4, 1]
		database = projectinfo.iloc[5, 1]

		target_path = sys.argv[2]
		types = sys.argv[4]
		os.chdir(target_path)

		sampleInfo = pd.read_csv('samples.txt', sep='\t', header=None)
		origi_record = pd.read_excel('原始记录.xlsx', sheet_name=1).fillna('')
		statistic = pd.read_csv('Evaluation/Statistic.csv')
		
		go_file = os.path.join('报告及附件', '3-3功能分析', '3-3-3GO功能分析', f'{groupvs}_GO', 'GO.xlsx')
		go = pd.read_excel(go_file, sheet_name=None)

		kegg_file = os.path.join('报告及附件', '3-3功能分析', '3-3-4KEGG通路分析', f'{groupvs}_KEGG', 'kegg.xlsx')
		kegg = pd.read_excel(kegg_file, sheet_name=None)

		protein_file = os.path.join('报告及附件', '3-1鉴定数量分析', '3-1-1 鉴定与定量结果统计', '附件1_蛋白质鉴定列表.xlsx')
		protein = pd.read_excel(protein_file, sheet_name=0)

		peptideScore = os.path.join('报告及附件', '5附件与说明文档', '5-2质量控制（QC）', 'PeptideScore.txt')
		with open(peptideScore) as f:
			medianScore = f.readline().split('=')[-1].strip()
			percentage = re.split('[=(]', f.readline())[1].strip()
	except Exception as e:
		print(e)
		# raise
		exit()
	
	def paragraph_format(self, pa, size, family, r = 0x00, g = 0x00, b = 0x00, bold = None):
		pa.font.size = Pt(size)
		pa.font.name = family
		if bold == True:
			pa.font.bold = True
		pa.font.color.rgb = RGBColor(r, g, b)
		pa._element.rPr.rFonts.set(qn('w:eastAsia'), family)

	def text_replace(self, p, text1_list, text2_list):
		'''
		p:段落
		text1_list: 需要替换的文本列表
		text2_list: 替换的文本列表,与上面列表长度相等
		'''
		text1_list = text1_list[0] if len(text1_list) == 1 else '|'.join(text1_list)
		text = p.text.strip()
		text_list = re.split(text1_list, text)
		p.clear()
		for i in range(len(text_list) - 1):
			self.paragraph_format(p.add_run(text_list[i]), size = 10.5, family = u'微软雅黑')
			self.paragraph_format(p.add_run(text2_list[i]), size = 10.5, family = 'Arial')
		self.paragraph_format(p.add_run(text_list[-1]), size = 10.5, family = u'微软雅黑')

	def insert_table(self, data, table, axis=0):
		'''
		data:插入的数据，格式为dataframe
		table:待插入数据的表格
		axis : 0为行添加，1为列添加
		'''
		total_row = len(table.rows)
		total_columns = len(table.columns)
		if axis == 0:
			row_cells = table.rows[total_row - 1].cells
			for row_num in range(data.shape[0]):		
				if row_num > 0:
					row_line = table.add_row()
					row_line.height = Cm(0.7)
					row_cells = row_line.cells
					# row_cells = table.rows[total_row + row_num -1].cells			
				for col_num in range(data.shape[1]):			
					tmp = data.iloc[row_num, col_num]
					if type(tmp) == 'float':
						tmp = int(tmp)
					pa = row_cells[col_num].paragraphs[0].add_run(str(tmp))
					row_cells[col_num].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					row_cells[col_num].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=9, family="Arial")
		else:
			if data.shape[0] > 2:
				for i in range(data.shape[0] - 2):
					row_line = table.add_row()
					row_line.height = Cm(1)
					row_cells = row_line.cells
					pa = row_cells[0].paragraphs[0].add_run('样本名称')
					row_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=9, family=u'微软雅黑')

			col_cells = table.columns[total_columns - 1].cells
			for col_num in range(data.shape[1]):
				if col_num > 0:
					col_line = table.add_column(Inches(0.7))
					col_cells = col_line.cells
					col_cells = table.columns[total_columns + col_num -1].cells
				for row_num in range(data.shape[0]):
					tmp = data.iloc[row_num, col_num]
					pa = col_cells[row_num].paragraphs[0].add_run(str(tmp))
					col_cells[row_num].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					col_cells[row_num].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=9, family="Arial")

	def insert_png(self, p, png1, png2):
		if os.path.exists(png2):
			p.text = p.text.strip().replace(png1, '')
			run = p.add_run()
			run.add_picture(png2, width=Inches(4.5))

	def extract_top(self, data, func, num=5):
		top = []
		if func == 'BP':
			data = data['Enrichment']
			data = data[data['Category'] == 'P'].head(num)
			top = data['Term'].tolist()
		elif func == 'MF':
			data = data['Enrichment']
			data = data[data['Category'] == 'F'].head(num)
			top = data['Term'].tolist()
		elif func == 'CC':
			data = data['Enrichment']
			data = data[data['Category'] == 'C'].head(num)
			top = data['Term'].tolist()
		elif func == 'goEnrich':
			data = data['Enrichment'].head(num)
			top = data['Term'].tolist()
		elif func == 'keggEnrich':
			data = data['Enrichment'].head(num)
			top = data['Map_Name'].tolist()
		elif func == 'keggID':
			data = data['Enrichment'].head(num)
			top = data['Map_ID'].tolist()
		elif func == 'map2query':
			data = data['map2query'].head(num)
			top = data['Map_Name'].tolist()
		return top

	def header(self, paragraphs, start_row=10):
		today = str(datetime.date.today())
		for i in range(3):
			if i in range(0, 2):
				pa = paragraphs[i + start_row].add_run(Template.projectinfo.iloc[i, 1])
				self.paragraph_format(pa, size=14, family=u'微软雅黑', bold=True)
			if i in range(2, 3):
				pa = paragraphs[i + start_row].add_run(Template.projectinfo.iloc[i, 1])
				self.paragraph_format(pa, size=14, family="Arial", bold=True) #### family = 'Calibri'
		pa = paragraphs[start_row + 3].add_run(today)
		self.paragraph_format(pa, size=14, family="Arial", bold=True)

	def record(self, record):		
		record_diff = record[record['group'].str.contains('vs|oneway')]
		if Template.types == 'l':
			columns = ['group', 'Decreased(%)', 'Increased(%)', 'all(%)', 'absence', 'presence']
			record_diff = record_diff[columns]
			record_diff['Decreased(%)'] = record_diff['Decreased(%)'].apply(lambda x: x.split('(')[0])
			record_diff['Increased(%)'] = record_diff['Increased(%)'].apply(lambda x: x.split('(')[0])
			record_diff['all(%)'] = record_diff['all(%)'].apply(lambda x: x.split('(')[0])
			record_diff['absence'] = record_diff['absence'].apply(lambda x: int(x))
			record_diff['presence'] = record_diff['presence'].apply(lambda x: int(x))
		else:
			columns = ['group', 'Decreased(%)', 'Increased(%)', 'all(%)']
			record_diff = record_diff[columns]
			record_diff['Decreased(%)'] = record_diff['Decreased(%)'].apply(lambda x: x.split('(')[0])
			record_diff['Increased(%)'] = record_diff['Increased(%)'].apply(lambda x: x.split('(')[0])
			record_diff['all(%)'] = record_diff['all(%)'].apply(lambda x: x.split('(')[0])
		return record_diff


	# =============================================================================
	# 插入表格数据
	# =============================================================================

	def table_data(self, tables, fc):
		summary1 = tables[0]
		self.paragraph_format(summary1.cell(2,0).paragraphs[0].add_run(Template.species), size = 9, family = 'Arial')
		self.paragraph_format(summary1.cell(2,1).paragraphs[0].add_run(Template.database), size = 9,family = 'Arial')
		
		# sampleInfo.iloc[:, 0] = sampleInfo.iloc[:, 0].apply(lambda x: re.sub('-\d+$', '', x))
		def ab(df):return';'.join(df.values)
		samples_data = pd.DataFrame(Template.sampleInfo.groupby(2)[1].apply(ab))  ## 多行合并一行
		samples_data.insert(0,'', samples_data.index.tolist())
		self.insert_table(samples_data, tables[1])
		
		statistic_list = [j for i in Template.statistic.values.T.tolist() for j in i]
		statistic_df = pd.DataFrame(statistic_list).T
		self.insert_table(statistic_df, tables[2])
		self.insert_table(statistic_df, tables[5])

		diff_table = tables[3]
		tmp_text = diff_table.cell(1, 1).text
		tmp_text = tmp_text.replace('upRatio', str(fc))
		diff_table.cell(1, 1).text = ''
		diff_table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		diff_table.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		self.paragraph_format(diff_table.cell(1, 1).paragraphs[0].add_run(tmp_text), size = 9, family = 'Arial')
		record_diff = self.record(Template.origi_record)
		self.insert_table(record_diff, tables[3])

		func_table = tables[4]
		self.paragraph_format(func_table.cell(2,0).paragraphs[0].add_run(Template.groupvs), size = 9, family = 'Arial')
		self.paragraph_format(func_table.cell(8,0).paragraphs[0].add_run(Template.groupvs), size = 9, family = 'Arial')		
		goEnrich_top5 = self.extract_top(Template.go, 'goEnrich')
		for i in range(5):
			self.paragraph_format(func_table.cell(i+2,1).paragraphs[0].add_run(goEnrich_top5[i]),size = 9, family ='Arial')		
		keggEnrich_top5 = self.extract_top(Template.kegg, 'keggEnrich')
		for i in range(5):
			self.paragraph_format(func_table.cell(i+8,1).paragraphs[0].add_run(keggEnrich_top5[i]),size = 9, family ='Arial')

				
		if Template.types == 'l':
			LFQ_list = [i for i in Template.protein.columns if re.search('LFQ', i)]
			sample = [i.split(' ')[-1] for i in LFQ_list]
			sample.insert(0, 'Total')
			proNum = [Template.protein[i].count() for i in LFQ_list]
			proNum.insert(0, Template.protein['Protein'].count())
			database_list = [Template.database] * len(sample)
			data_frame_dict = {'database': database_list, 'sample': sample, 'proNum': proNum}
			data_frame = pd.DataFrame(data_frame_dict)
			self.insert_table(data_frame, tables[6])
			self.insert_table(record_diff, tables[7])
			pa = tables[8].cell(8,1).paragraphs[0].add_run(Template.database)
			self.paragraph_format(pa, size=9, family="Arial")
		else:
			self.insert_table(record_diff, tables[6])
			pa = tables[8].cell(7,1).paragraphs[0].add_run(Template.database)
			self.paragraph_format(pa, size=9, family="Arial")

			itraq_tmt = Template.sampleInfo.T.iloc[:2,:]
			itraq_tmt.iloc[0, :] = itraq_tmt.iloc[0, :].apply(lambda x: x.split('.')[-1])
			if len(itraq_tmt.iloc[0,:]) != len(set(itraq_tmt.iloc[0, :])):
				newdf = pd.DataFrame()
				for i, j in itraq_tmt.T.groupby(0):
					newdf[i] = j[1].tolist()
				newdf = newdf.T
				newdf.insert(0, '', newdf.index.tolist())
				newdf = newdf.T
				newdf.reset_index(inplace=True)
				newdf.drop('index', axis =1, inplace=True)
				No = newdf.index.tolist()
				No[0] = 'No.'
				newdf['No'] = No
				self.insert_table(newdf, tables[7], axis=1)
			else:
				No = itraq_tmt.index.tolist()
				No[0] = 'No.'
				itraq_tmt['No'] = No
				self.insert_table(itraq_tmt, tables[7], axis=1)

	# =============================================================================
	# 插入文本及结果图片
	# =============================================================================
	def text_png_data(self, paragraphs, fc):
		venn2 = ''
		if os.path.exists(os.path.join('报告及附件', '3-1鉴定数量分析', '3-1-3 韦恩图分析')):
			venn2 = [i for i in os.listdir(os.path.join('报告及附件', '3-1鉴定数量分析', '3-1-3 韦恩图分析', '组间')) if re.search('png', i)][0]
		keggID = self.extract_top(Template.kegg, 'keggID')[0]
		bp_top5 = self.extract_top(Template.go, 'BP')
		mf_top5 = self.extract_top(Template.go, 'MF')
		cc_top5 = self.extract_top(Template.go, 'CC')
		goEnrich_top5 = self.extract_top(Template.go, 'goEnrich')
		keggEnrich_top5 = self.extract_top(Template.kegg, 'keggEnrich')
		map2query_top5 = self.extract_top(Template.kegg, 'map2query')
		groupNum = str(len(set(Template.sampleInfo.iloc[:,2])))
		Num = list(Counter(Template.sampleInfo.iloc[:,2]).values())
		Num = ','.join(list(set(Num))) if len(set(Num)) > 1 else str(Num[0])
		total_num = str(len(Template.sampleInfo.iloc[:,2]))
		for p in paragraphs:
			if 'upRatio' in p.text:
				find_num = len(re.findall('upRatio', p.text))
				self.text_replace(p, ['upRatio'] * find_num, [str(fc)] * find_num)
			if 'downRatio' in p.text:
				find_num = len(re.findall('downRatio', p.text))
				self.text_replace(p, ['downRatio'] * find_num, [str(round(1 / fc, 2))] * find_num)
			if 'groupvs' in p.text:
				self.text_replace(p, ['groupvs'], [Template.groupvs])
			if 'BP-TOP5' in p.text:
				self.text_replace(p, ['BP-TOP5', 'MF-TOP5', 'CC-TOP5'], [bp_top5, mf_top5, cc_top5])
			if 'kegg-map2query-top5' in p.text:
				self.text_replace(p, ['kegg-map2query-top5'], ['，'.join(map2query_top5)])
			if 'KeggEnrich-top5' in p.text:
				self.text_replace(p, ['KeggEnrich-top5', 'KeggEnrich-top1'], ['，'.join(keggEnrich_top5), keggEnrich_top5[0]])
			if 'Percentage' in p.text:
				self.text_replace(p, ['Percentage', 'Median Score'], [Template.percentage, Template.medianScore])
			if 'groupNum' in p.text:
				self.text_replace(p, ['groupNum', 'Num', 'total'], [groupNum, Num, total_num])

			if '[Statistic]' in p.text:
				self.insert_png(p, '[Statistic]', os.path.join('报告及附件', '3-1鉴定数量分析', '3-1-1 鉴定与定量结果统计', '鉴定与定量结果统计柱状图.png'))
			if '[venn1]' in p.text:
				self.insert_png(p, '[venn1]', os.path.join('报告及附件', '3-1鉴定数量分析', '3-1-3 韦恩图分析', '组内', f'venn_{Template.groupvs.split("_vs_")[0]}.png'))
			if venn2:
				if '[venn2]' in p.text:
					self.insert_png(p, '[venn2]', os.path.join('报告及附件', '3-1鉴定数量分析', '3-1-3 韦恩图分析', '组间', f'{venn2}'))
			if '[pro_diff]' in p.text:
				self.insert_png(p, '[pro_diff]', os.path.join('报告及附件', '3-2表达差异分析', '3-2-1差异结果数量统计', '蛋白质定量差异结果柱状图.png'))
			if '[volcano]' in p.text:
				self.insert_png(p, '[volcano]', os.path.join('报告及附件', '3-2表达差异分析', '3-2-2火山图', f'Volcano_Plot_{Template.groupvs}.png'))
			if '[cluster]' in p.text:
				self.insert_png(p, '[cluster]', os.path.join('报告及附件', '3-2表达差异分析', '3-2-3聚类分析', f'{Template.groupvs}_Cluster', 'cluster1.png'))
			if '[Subcellular_Localization]' in p.text:
				self.insert_png(p, '[Subcellular_Localization]', os.path.join('报告及附件', '3-3功能分析', '3-3-1亚细胞定位分析', f'{Template.groupvs}_Cello', 'Subcellular_Localization.png'))
			if '[TopDomainStat]' in p.text:
				self.insert_png(p, '[TopDomainStat]', os.path.join('报告及附件', '3-3功能分析', '3-3-2结构域分析', f'{Template.groupvs}_Domain', 'TopDomainStat.png'))
			if '[Domain_Enrichment]' in p.text:
				self.insert_png(p, '[Domain_Enrichment]', os.path.join('报告及附件', '3-3功能分析', '3-3-2结构域分析', f'{Template.groupvs}_Domain', 'Domain_Enrichment.png'))
			if '[GOLevel2]' in p.text:
				self.insert_png(p, '[GOLevel2]', os.path.join('报告及附件', '3-3功能分析', '3-3-3GO功能分析', f'{Template.groupvs}_GO', 'GOLevel2.png'))
			if '[BP_Enrichment]' in p.text:
				self.insert_png(p, '[BP_Enrichment]', os.path.join('报告及附件', '3-3功能分析', '3-3-3GO功能分析', f'{Template.groupvs}_GO', 'BP_Enrichment.png'))
			if '[CC_Enrichment]' in p.text:
				self.insert_png(p, '[CC_Enrichment]', os.path.join('报告及附件', '3-3功能分析', '3-3-3GO功能分析', f'{Template.groupvs}_GO', 'CC_Enrichment.png'))
			if '[MF_Enrichment]' in p.text:
				self.insert_png(p, '[MF_Enrichment]', os.path.join('报告及附件', '3-3功能分析', '3-3-3GO功能分析', f'{Template.groupvs}_GO', 'MF_Enrichment.png'))
			if '[BP_DAG]' in p.text:
				self.insert_png(p, '[BP_DAG]', os.path.join('报告及附件', '3-3功能分析', '3-3-3GO功能分析', f'{Template.groupvs}_GO', 'BP_DAG.png'))
			if '[CC_DAG]' in p.text:
				self.insert_png(p, '[CC_DAG]', os.path.join('报告及附件', '3-3功能分析', '3-3-3GO功能分析', f'{Template.groupvs}_GO', 'CC_DAG.png'))
			if '[MF_DAG]' in p.text:
				self.insert_png(p, '[MF_DAG]', os.path.join('报告及附件', '3-3功能分析', '3-3-3GO功能分析', f'{Template.groupvs}_GO', 'MF_DAG.png'))
			if '[TopMapStat]' in p.text:
				self.insert_png(p, '[TopMapStat]', os.path.join('报告及附件', '3-3功能分析', '3-3-4KEGG通路分析', f'{Template.groupvs}_KEGG', 'TopMapStat.png'))
			if '[KEGG_Enrichment]' in p.text:
				self.insert_png(p, '[KEGG_Enrichment]', os.path.join('报告及附件', '3-3功能分析', '3-3-4KEGG通路分析', f'{Template.groupvs}_KEGG', 'KEGG_Enrichment.png'))
			if '[kegg_pathway]' in p.text:
				self.insert_png(p, '[kegg_pathway]', os.path.join('报告及附件', '3-3功能分析', '3-3-4KEGG通路分析', f'{Template.groupvs}_KEGG', 'map', f'{keggID}.png'))
			if '[ppi]' in p.text:
				self.insert_png(p, '[ppi]', os.path.join('报告及附件', '3-3功能分析', '3-3-5蛋白互作网络分析', f'{Template.groupvs}_PPI', 'ppi.png'))
			if '[Module_ppi]' in p.text:
				self.insert_png(p, '[Module_ppi]', os.path.join('报告及附件', '3-3功能分析', '3-3-5蛋白互作网络分析', f'{Template.groupvs}_PPI', 'Module_ppi.png'))
			if '[mass_error]' in p.text:
				self.insert_png(p, '[mass_error]', os.path.join('报告及附件', '5附件与说明文档', '5-2质量控制（QC）', 'mass_error_distribution.png'))
			if '[Andromeda_Score_Distribution]' in p.text:
				self.insert_png(p, '[Andromeda_Score_Distribution]', os.path.join('报告及附件', '5附件与说明文档', '5-2质量控制（QC）', 'Andromeda_Score_Distribution.png'))
				self.insert_png(p, '[Andromeda_Score_Distribution]', os.path.join('报告及附件', '5附件与说明文档', '5-2质量控制（QC）', 'Ion_Score_Distribution.png'))
			if '[pI_Distribution]' in p.text:
				self.insert_png(p, '[pI_Distribution]', os.path.join('Evaluation', 'pI_Distribution.png'))
			if '[Ratio_Distribution]' in p.text:
				self.insert_png(p, '[Ratio_Distribution]', os.path.join('报告及附件', '5附件与说明文档', '5-2质量控制（QC）', f'Ratio_Distribution_{Template.groupvs}.png'))
			if '[MW_Distribution]' in p.text:
				self.insert_png(p, '[MW_Distribution]', os.path.join('报告及附件', '5附件与说明文档', '5-3鉴定蛋白及肽段特性描述', 'Molecular_Weight_Distribution.png'))
			if '[PepCount_Distribution]' in p.text:
				self.insert_png(p, '[PepCount_Distribution]', os.path.join('报告及附件', '5附件与说明文档', '5-3鉴定蛋白及肽段特性描述', 'Peptide_Length_Distribution.png'))
			if '[Protein_Sequence_Coverage_Distribution]' in p.text:
				self.insert_png(p, '[Protein_Sequence_Coverage_Distribution]', os.path.join('报告及附件', '5附件与说明文档', '5-3鉴定蛋白及肽段特性描述', 'Protein_Sequence_Coverage_Distribution.png'))
			if '[Peptide_Count_Distribution]' in p.text:
				self.insert_png(p, '[Peptide_Count_Distribution]', os.path.join('报告及附件', '5附件与说明文档', '5-3鉴定蛋白及肽段特性描述', 'Peptide_Count_Distribution.png'))


	def save(self):
		pass

