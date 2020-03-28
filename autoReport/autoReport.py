import os
import sys
import click
from docx import Document

@click.command()
@click.option("--input_path", '-i', required=True, help="项目路径")
@click.option("--types", '-t', required=True, help="项目类型, l或L为labfree; i或I为Itraq; t或T为TMT")
@click.option('--fc', '-fc', help='FoldChange')

def main(input_path, types, fc):
	types = types.lower()
	types_dict = {'l': 'Labelfree报告模板-20200309.docx',
				  'i': 'iTRAQ报告模板-20200325-4标8标.docx',
				  't': 'TMT报告模板-20200325-6标10标16标.docx'}

	template_file = os.path.join(os.getcwd(), 'template', types_dict.get(types))
	document = Document(template_file)
	paragraphs = document.paragraphs
	tables = document.tables

	if types == 'l':	
		fc = 2 if fc == None else float(fc)
		from templatePy.project import Labfree
		labfree = Labfree()
		labfree.header(paragraphs)
		labfree.table_data(tables, fc)
		labfree.text_png_data(paragraphs, fc)
		labfree.save(document)
	elif types == 'i':
		fc = 1.2 if fc == None else float(fc)
		from templatePy.project import Itraq
		itraq = Itraq()
		itraq.header(paragraphs, start_row=13)
		itraq.table_data(tables, fc)
		itraq.text_png_data(paragraphs, fc)
		itraq.save(document)
	elif types == 't':
		fc = 1.2 if fc == None else float(fc)
		from templatePy.project import TMT
		tmt = TMT()
		tmt.header(paragraphs, start_row=13)
		tmt.table_data(tables, fc)
		tmt.text_png_data(paragraphs, fc)
		tmt.save(document)



if __name__ == '__main__':
	main()
	

