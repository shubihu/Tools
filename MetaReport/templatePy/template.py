# -*- coding: future_fstrings -*-     # should work even without -*-
###
import os
import sys
import re
import datetime
import pandas as pd
pd.options.mode.chained_assignment = None
# import modin.pandas as pd
from collections import Counter
from collections import OrderedDict
# from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
# import win32com.client

def extract_top(data, func, num=5):
		top = []
		pathway = []
		if func == 'BP':
			data = data['Enrichment']
			data = data[data['Category'] == 'P'].head(num)
			top = data['Term'].tolist()		
		elif func == 'MF':
			data = data['Enrichment']
			data = data[data['Category'] == 'F'].head(num)
			top = data['Term'].tolist()
		elif func == 'CC':
			data = data['Enrichment']
			data = data[data['Category'] == 'C'].head(num)
			top = data['Term'].tolist()
		elif func == 'goEnrich':
			data = data['Enrichment'].head(num)
			top = data['Term'].tolist()
		elif func == 'keggEnrich':
			data = data['Enrichment'].head(num)
			top = data['Map_Name'].tolist()
			pathway = top[0]
		elif func == 'keggID':
			data = data['Enrichment'].head(num)
			top = data['Map_ID'].tolist()
		elif func == 'map2query':
			data = data['map2query'].head(num)
			top = data['Map_Name'].tolist()

		if not func in ['keggID', 'map2query']:
			if 'P value' in data.columns:
				pvalue_list = data['P value'].tolist()
			elif 'p.value' in data.columns:
				pvalue_list = data['p.value'].tolist()
			pvalue_top = [i for i in pvalue_list if float(i) < 0.05]
			top = top[:len(pvalue_top)] if pvalue_top else []
		return top, pathway


class Template:
	def __init__(self):
		pass
	
	def paragraph_format(self, pa, size, family, r = 0x00, g = 0x00, b = 0x00, bold = None):
		pa.font.size = Pt(size)
		pa.font.name = family
		if bold == True:
			pa.font.bold = True
		pa.font.color.rgb = RGBColor(r, g, b)
		pa._element.rPr.rFonts.set(qn('w:eastAsia'), family)

	def delete_paragraph(self, paragraphs, p_index_list):
		'''
		p_index_list:需要删除段落索引的列表
		'''
		for i in p_index_list:
			p = paragraphs[i]
			p = p._element
			if p.getparent() is not None:
				p.getparent().remove(p)
				p._p = p._element = None

	def text_replace(self, p, text1_list, text2_list, size=10.5, family_ch=u'微软雅黑', family_en='Arial', bold=None):
		'''
		p:段落
		text1_list: 需要替换的文本列表
		text2_list: 替换的文本列表,与上面列表长度相等
		'''
		text1_list = [i.strip('[]') for i in text1_list]
		text1_list = text1_list[0] if len(text1_list) == 1 else '|'.join(text1_list)
		text = p.text.strip()
		text_list = re.split(text1_list, text)
		p.clear()
		for i in range(len(text_list) - 1):
			self.paragraph_format(p.add_run(text_list[i]), size = size, family = family_ch, bold=bold)
			if '无P值小于0.05' in text2_list[i]:
				self.paragraph_format(p.add_run(text2_list[i]), size = size, family = family_ch, bold=bold)
			else:
				self.paragraph_format(p.add_run(text2_list[i]), size = size, family = family_en, bold=bold)
		self.paragraph_format(p.add_run(text_list[-1]), size = size, family = family_ch, bold=bold)

	def Set_Background_Color(self, cell, rgbColor):
		'''
		#定义单元格填充颜色函数
		'''
		shading_elm = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = rgbColor)) #固定写法，照抄即可
		cell._tc.get_or_add_tcPr().append(shading_elm)

	def insert_table(self, data, table, axis=0, size=9, family_ch=u'微软雅黑', family_en='Arial', bold=None, rgbColor=None):
		'''
		data:插入的数据，格式为dataframe
		table:待插入数据的表格
		axis : 0为行添加，1为列添加
		'''
		total_row = len(table.rows)
		total_columns = len(table.columns)
		if axis == 0:
			row_cells = table.rows[total_row - 1].cells
			for row_num in range(data.shape[0]):
				if row_num > 0:
					row_line = table.add_row()
					row_line.height = Cm(0.7)
					row_cells = row_line.cells
					# row_cells = table.rows[total_row + row_num -1].cells			
				for col_num in range(data.shape[1]):
					tmp = data.iloc[row_num, col_num]
					if type(tmp) == 'float':
						tmp = int(tmp)
					if rgbColor:
						self.Set_Background_Color(row_cells[col_num], rgbColor)
					pa = row_cells[col_num].paragraphs[0].add_run(str(tmp))
					row_cells[col_num].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					row_cells[col_num].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=size, family=family_en, bold=bold)
		else:
			if data.shape[0] > 2:
				for i in range(data.shape[0] - 2):
					row_line = table.add_row()
					row_line.height = Cm(1)
					row_cells = row_line.cells
					pa = row_cells[0].paragraphs[0].add_run('样本名称')
					row_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=size, family=family_ch, bold=bold)

			col_cells = table.columns[total_columns - 1].cells
			for col_num in range(data.shape[1]):
				if col_num > 0:
					col_line = table.add_column(Inches(0.7))
					col_cells = col_line.cells
					col_cells = table.columns[total_columns + col_num -1].cells
				for row_num in range(data.shape[0]):
					tmp = data.iloc[row_num, col_num]
					pa = col_cells[row_num].paragraphs[0].add_run(str(tmp))
					col_cells[row_num].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					col_cells[row_num].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=size, family=family_en, bold=bold)

	def table_center(self, table):
		'''
		table 表格数据居中
		'''
		for i in range(len(table.rows)):
			for j in range(len(table.columns)):
				table.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
				table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# return table

	def move_table_after(self, table, paragraph):
		'''
		移动表格到指定位置
		'''
		tbl, p = table._tbl, paragraph._p
		p.addnext(tbl)

	def insert_png(self, p, png1, png2, png3=None):
		if os.path.exists(png2):
			p.text = p.text.strip().replace(png1, '')
			run = p.add_run()
			if png3:
				run.add_picture(png2, width=Inches(2.5))
				run.add_picture(png3, width=Inches(2.5))
			else:
				run.add_picture(png2, width=Inches(4.5))
			return True
		else:
			png1_bak = png1.replace('[', '').replace(']', '')
			self.text_replace(p, [png1], [f'no {png1_bak}'])
			return False

	def record(self, record):		
		record_diff = record[record.loc[:, 'group'].str.contains('vs|oneway|twoway')]
		tmp = record_diff.iloc[:, 1:].applymap(lambda x: re.sub(r'\(.*\)', '', str(x)))  # applymap对每个元素进行处理
		tmp = tmp.applymap(lambda x: x if '-' in str(x) else int(float(x)))
		record_diff = pd.concat([record_diff.iloc[:, 0], tmp], axis=1)
		return record_diff

	def header(self, paragraphs, start_row=10):
		pass

	# =============================================================================
	# 插入表格数据
	# =============================================================================
	def table_data(self, tables):
		pass

	# =============================================================================
	# 插入文本及结果图片
	# =============================================================================
	def text_png_data(self, document, paragraphs):
		pass


	def save(self):
		pass

