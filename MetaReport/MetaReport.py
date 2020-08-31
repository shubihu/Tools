

import os
import sys
import click
from docx import Document
from templatePy.update import get_version
from templatePy.update import update

version = 1.04
newversion = get_version()

if version < newversion:
	print('服务器上发现最新版本，开始下载更新代码。\n服务器(192.168.130.252)上代码地址：/database/metabolome/code_Version2/MetaReport')
	update(os.path.abspath(os.path.dirname(sys.argv[0])))
	exit()

def lipids(*args):
	input_path, _, document, paragraphs, tables = args
	from templatePy.lipids import Lipids
	lipids = Lipids(input_path)
	lipids.header(paragraphs)
	lipids.table_data(tables)
	lipids.text_png_data(paragraphs)
	lipids.save(document)

def fattyacid(*args):
	input_path, types, document, paragraphs, tables = args
	from templatePy.fattyacid import Fattyacid
	fa = Fattyacid(input_path, types)
	fa.header(paragraphs)
	fa.table_data(tables)
	fa.text_png_data(paragraphs)
	fa.save(document)

def notarget(*args):
	input_path, _, document, paragraphs, tables = args
	from templatePy.notarget import Notarget
	nt = Notarget(input_path)
	nt.header(paragraphs)
	nt.table_data(tables)
	nt.text_png_data(paragraphs)
	nt.save(document)

@click.command()
@click.option("--input_path", '-i', required=True, help="项目路径")
@click.option("--types", '-t', required=True, help="项目类型, ld为lipids, lc为长链脂肪酸，sc为短链脂肪酸, nt为非靶代谢")
def main(input_path, types):
	types = types.lower()
	types_dict = {'ld': '高分辨广谱脂质组绝对定量报告模板-FINAL.docx',
				  'lc': '中长链脂肪酸报告模板.docx',
				  'sc': '短链脂肪酸报告模板.docx',
				  'nt': '非靶代谢报告模板.docx'}

	project_dict = {'ld': lipids, 'lc': fattyacid, 'sc': fattyacid, 'nt': notarget}

	template_file = os.path.join(os.path.abspath(os.path.dirname(sys.argv[0])), 'template', types_dict.get(types))
	document = Document(template_file)
	paragraphs = document.paragraphs
	tables = document.tables

	project_dict.get(types)(input_path, types, document, paragraphs, tables)

if __name__ == '__main__':
	main()
	

