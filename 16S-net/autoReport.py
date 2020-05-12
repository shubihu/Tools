import os
import sys
import shutil
import click
from docx import Document

platform = sys.platform

def s16(*args):
	input_path, types, document, paragraphs, tables = args
	from templatePy.template import Template
	s16 = Template(input_path, types)
	s16.header(paragraphs)
	s16.table_data(tables)
	s16.text_png_data(paragraphs)
	s16.save(document)
	s16.netReport()


@click.command()
@click.option("--input_path", '-i', required=True, help="项目路径")
@click.option("--types", '-t', required=True, help="项目类型, s为16S")
def main(input_path, types):
	types = types.lower()
	types_dict = {'s': '16S-ReportDemo.docx'}
	project_dict = {'s': s16}

	if os.path.exists(os.path.join(input_path, 'Result', 'WebReport')):
		shutil.rmtree(os.path.join(input_path, 'Result', 'WebReport'))
	shutil.copytree(os.path.join(os.path.abspath(os.path.dirname(sys.argv[0])), 'templateNet'), os.path.join(input_path, 'Result', 'WebReport'))

	template_file = os.path.join(os.path.abspath(os.path.dirname(sys.argv[0])), 'template', types_dict.get(types))
	document = Document(template_file)
	paragraphs = document.paragraphs
	tables = document.tables

	project_dict.get(types)(input_path, types, document, paragraphs, tables)

if __name__ == '__main__':
	main()
	

