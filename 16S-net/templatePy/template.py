import os
import sys
import re
import glob
import datetime
import pandas as pd
pd.options.mode.chained_assignment = None
# import modin.pandas as pd
from collections import Counter
# from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
# import win32com.client


class Template:
	def __init__(self, path, types):
		self.types = types
		os.chdir(path)
		# try :
		information_file = [i for i in os.listdir('.') if re.search('infor?mation', i, re.I)]
		if information_file:
			self.projectinfo = pd.read_excel(information_file[0], header=None)
		else:
			print('Error:该项目下无project_information表')
			exit()
		self.project_name = self.projectinfo.iloc[0, 1]
		self.school = self.projectinfo.iloc[1, 1]
		self.project_num = self.projectinfo.iloc[2, 1]
		self.today = str(datetime.date.today())
		self.groupvs = self.projectinfo.iloc[4, 1]

		sample_info = pd.read_csv('sample_info.txt', sep='\t')
		sample_info_dict = {i: len(j) for i, j in sample_info.groupby('Group').groups.items()}
		self.sample_info_df = pd.DataFrame.from_dict(sample_info_dict, orient='index')
		self.sample_info_df.insert(0, 'group', self.sample_info_df.index.tolist())
		self.sample_info_df['状态'] = ['固体'] * self.sample_info_df.shape[0]

		self.total_sample = str(sample_info.shape[0])
		self.group_num = str(self.sample_info_df.shape[0])
		sample_num = [str(i) for i in self.sample_info_df.iloc[:, 1]]
		self.sample_num = str(sample_num[0]) if len(set(sample_num)) == 0 else '/'.join(set(sample_num))

		self.qc = pd.read_csv('Result/01_QC/QcStatic.csv').head(5)
		otu_table = pd.read_csv('Result/02_OTU_Taxa/otu_table_tax.csv').head(5)
		self.otu_table = otu_table.iloc[:, [0, 1, 2, 3, -1]]

		self.sample_tree = [i for i in os.listdir("Result/03_Community/taxa_tree/sample_tree/") if re.search('png', i)][0]

		self.alpha_index = pd.read_csv('Result/04_Alpha_diversity/alpha_diversity_index.csv').head(5)
		tmp = self.alpha_index.iloc[:, 1:].applymap(lambda x: str(round(float(x), 3)))
		self.alpha_index = pd.concat([self.alpha_index.iloc[:, :1], tmp], axis=1)
		self.alpha_index['observed_species'] = self.alpha_index['observed_species'].apply(lambda x: str(int(float(x))))

		anosim_file = [file for file in glob.glob('Result/05_Beta_diversity/anosim/*.csv')]
		adonis_file = [file for file in glob.glob('Result/05_Beta_diversity/adonis/*.csv')]
		self.anosim = pd.read_csv(anosim_file[0])
		self.anosim['R'] = self.anosim['R'].apply(lambda x: round(x, 3))
		self.adonis = pd.read_csv(adonis_file[0])
		self.adonis['R2'] = self.adonis['R2'].apply(lambda x: round(x, 3))

		self.ko = pd.read_csv('Result/07_FunctionPrediction/KEGG/ko_prediction.csv', sep=',')
		self.ko = self.ko.iloc[0:5, 0:5]

		self.cog = pd.read_csv('Result/07_FunctionPrediction/COG/cog_prediction.csv', sep=',')
		self.cog = self.cog.iloc[0:5:, 0:5]

		# except Exception as e:
		# 	print(e)
		# 	print('请检查project_information中信息是否填写正确')
		# 	# raise
		# 	exit()
	
	def paragraph_format(self, pa, size, family, r = 0x00, g = 0x00, b = 0x00, bold = None):
		pa.font.size = Pt(size)
		pa.font.name = family
		if bold == True:
			pa.font.bold = True
		pa.font.color.rgb = RGBColor(r, g, b)
		pa._element.rPr.rFonts.set(qn('w:eastAsia'), family)

	def delete_paragraph(self, paragraphs, p_index_list):
		'''
		p_index_list:需要删除段落索引的列表
		'''
		for i in p_index_list:
			p = paragraphs[i]
			p = p._element
			if p.getparent() != None:
				p.getparent().remove(p)
				p._p = p._element = None

	def text_replace(self, p, text1_list, text2_list, size=10.5, family_ch=u'微软雅黑', family_en='Arial', bold=None):
		'''
		p:段落
		text1_list: 需要替换的文本列表
		text2_list: 替换的文本列表,与上面列表长度相等
		'''
		text1_list = [i.strip('[]') for i in text1_list]
		text1_list = text1_list[0] if len(text1_list) == 1 else '|'.join(text1_list)
		text = p.text.strip()
		text_list = re.split(text1_list, text)
		p.clear()
		for i in range(len(text_list) - 1):
			self.paragraph_format(p.add_run(text_list[i]), size = size, family = family_ch, bold=bold)
			if 'P值小于0.05' in text2_list[i]:
				self.paragraph_format(p.add_run(text2_list[i]), size = size, family = family_ch, bold=bold)
			else:
				self.paragraph_format(p.add_run(text2_list[i]), size = size, family = family_en, bold=bold)
		self.paragraph_format(p.add_run(text_list[-1]), size = size, family = family_ch, bold=bold)

	def replace_table_header(self, data, table, r = 0x00, g = 0x00, b = 0x00, bold=None):
		'''
		data:替换的数据title
		table:需要替换的表
		'''
		for i in range(4):
			table.cell(0, i + 1).text = ''
			p = table.cell(0, i + 1).paragraphs[0].add_run(data.columns[i + 1])
			table.cell(0, i + 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
			table.cell(0, i + 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			self.paragraph_format(p, size=9, r=r, g=g, b=b, family="Times New Roman", bold=bold)

	def insert_table(self, data, table, axis=0, size=9, family_ch=u'微软雅黑', family_en='Arial'):
		'''
		data:插入的数据，格式为dataframe
		table:待插入数据的表格
		axis : 0为行添加，1为列添加
		'''
		total_row = len(table.rows)
		total_columns = len(table.columns)
		if axis == 0:
			row_cells = table.rows[total_row - 1].cells
			for row_num in range(data.shape[0]):
				if row_num > 0:
					row_line = table.add_row()
					row_line.height = Cm(0.7)
					row_cells = row_line.cells
					# row_cells = table.rows[total_row + row_num -1].cells
				for col_num in range(data.shape[1]):
					tmp = data.iloc[row_num, col_num]
					if type(tmp) == 'float':
						tmp = int(tmp)
					if type(tmp) != str and tmp.dtype == 'float64':
						tmp = round(tmp, 3)
					pa = row_cells[col_num].paragraphs[0].add_run(str(tmp))
					row_cells[col_num].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					row_cells[col_num].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=size, family=family_en)
		else:
			if data.shape[0] > 2:
				for i in range(data.shape[0] - 2):
					row_line = table.add_row()
					row_line.height = Cm(1)
					row_cells = row_line.cells
					pa = row_cells[0].paragraphs[0].add_run('样本名称')
					row_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=size, family=family_ch)

			col_cells = table.columns[total_columns - 1].cells
			for col_num in range(data.shape[1]):
				if col_num > 0:
					col_line = table.add_column(Inches(0.7))
					col_cells = col_line.cells
					col_cells = table.columns[total_columns + col_num -1].cells
				for row_num in range(data.shape[0]):
					tmp = data.iloc[row_num, col_num]
					pa = col_cells[row_num].paragraphs[0].add_run(str(tmp))
					col_cells[row_num].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					col_cells[row_num].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=size, family=family_en)

	def table_center(self, table):
		'''
		table 表格数据居中
		'''
		for i in range(len(table.rows)):
			for j in range(len(table.columns)):
				table.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
				table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# return table

	def move_table_after(self, table, paragraph):
		'''
		移动表格到指定位置
		'''
		tbl, p = table._tbl, paragraph._p
		p.addnext(tbl)

	def insert_png(self, p, png1, png2, png3=None):
		if os.path.exists(png2):
			p.text = p.text.strip().replace(png1, '')
			run = p.add_run()
			if png3:
				run.add_picture(png2, width=Inches(2.5))
				run.add_picture(png3, width=Inches(2.5))
			else:
				run.add_picture(png2, width=Inches(4.5))
			return True
		else:
			return False

	def header(self, paragraphs, start_row=10):
		for i in range(3):
			pa = paragraphs[i+5].add_run(self.projectinfo.iloc[i, 1])
			self.paragraph_format(pa, size=16, family=u'微软雅黑')
		pa = paragraphs[start_row - 2].add_run(self.today)
		self.paragraph_format(pa, size=14, family=u'微软雅黑')

	# =============================================================================
	# 插入表格数据
	# =============================================================================

	def table_data(self, tables):
		self.insert_table(self.sample_info_df, tables[0], family_en='Times New Roman')
		self.insert_table(self.qc, tables[1], family_en='Times New Roman')
		self.replace_table_header(self.otu_table, tables[2], 255, 255, 255, bold=True)
		self.insert_table(self.otu_table, tables[2], family_en='Times New Roman')
		self.insert_table(self.alpha_index, tables[3], family_en='Times New Roman')
		self.insert_table(self.anosim, tables[4], family_en='Times New Roman')
		self.insert_table(self.adonis, tables[5], family_en='Times New Roman')
		self.replace_table_header(self.ko, tables[6])
		self.insert_table(self.ko, tables[6], family_en='Times New Roman')
		self.replace_table_header(self.cog, tables[7])
		self.insert_table(self.cog, tables[7], family_en='Times New Roman')

	# =============================================================================
	# 插入文本及结果图片
	# =============================================================================
	def text_png_data(self, paragraphs):
		for p in paragraphs:
			if 'total_sample' in p.text:
				self.text_replace(p, ['total_sample', 'group_num', 'sample_num'], [self.total_sample, self.group_num, self.sample_num], family_ch=u'微软雅黑', family_en='Times New Roman')

			if "[venn.png]" in p.text:
				self.insert_png(p, "[venn.png]", f"Result/02_OTU_Taxa/Venn/{self.groupvs}/venn.png")

			if "[sample_tree.png]" in p.text:
				self.insert_png(p, "[sample_tree.png]", f"Result/03_Community/taxa_tree/sample_tree/{self.sample_tree}")

			if "[group_tree.png]" in p.text:
				self.insert_png(p, "[group_tree.png]", "Result/03_Community/taxa_tree/group_tree/level_tree.png")

			if "[sample_barplot.png]" in p.text:
				self.insert_png(p, "[sample_barplot.png]", "Result/03_Community/community/Top10_sample/Phylum/Phylum_top10.png")

			if "[group_barplot.png]" in p.text:
				self.insert_png(p, "[group_barplot.png]", "Result/03_Community/community/Top10_group/Phylum/Phylum_group_top10.png")

			if "[weight_upgma.png]" in p.text:
				self.insert_png(p, "[weight_upgma.png]", "Result/03_Community/UPGMA/weighted_unifrac/weighted_unifrac_dm_tree_barplot.png")

			if "[unweight_upgma.png]" in p.text:
				self.insert_png(p, "[unweight_upgma.png]", "Result/03_Community/UPGMA/unweighted_unifrac/unweighted_unifrac_dm_tree_barplot.png")

			if "[sample_heatmap.png]" in p.text:
				self.insert_png(p, "[sample_heatmap.png]", "Result/03_Community/taxa_heatmap/cluster_sample/Phylum.png")

			if "[group_heatmap.png]" in p.text:
				self.insert_png(p, "[group_heatmap.png]", "Result/03_Community/taxa_heatmap/cluster_group/Phylum_group.png")

			if "[group_heatmap.png]" in p.text:
				self.insert_png(p, "[group_heatmap.png]", "Result/03_Community/taxa_heatmap/cluster_group/Phylum_group.png")

			if "[rarefaction.png]" in p.text:
				self.insert_png(p, "[rarefaction.png]", "Result/04_Alpha_diversity/Rarefaction_Curve/observed_species_sample.png")

			if "[shannon.png]" in p.text:
				self.insert_png(p, "[shannon.png]", "Result/04_Alpha_diversity/Shannon/shannon_sample.png")

			if "[rank.png]" in p.text:
				self.insert_png(p, "[rank.png]", "Result/04_Alpha_diversity/Rank_Abundance/rank_sampleID.png")

			if "[specaccum.png]" in p.text:
				self.insert_png(p, "[specaccum.png]", "Result/04_Alpha_diversity/Specaccum/specaccum.png")

			if "[alpha_diff.png]" in p.text:
				self.insert_png(p, "[alpha_diff.png]", "Result/04_Alpha_diversity/Alpha_div_diff/observed_species/observed_species.png")

			if "[weight_beta_diff.png]" in p.text:
				self.insert_png(p, "[weight_beta_diff.png]", "Result/05_Beta_diversity/Beta_div_diff/weighted_unifrac/weighted_unifrac.png")

			if "[unweight_beta_diff.png]" in p.text:
				self.insert_png(p, "[unweight_beta_diff.png]", "Result/05_Beta_diversity/Beta_div_diff/unweighted_unifrac/unweighted_unifrac.png")

			if "[pca.png]" in p.text:
				self.insert_png(p, "[pca.png]", f"Result/05_Beta_diversity/PCA/{self.groupvs}/pca.png")

			if "[weight_pcoa.png]" in p.text:
				self.insert_png(p, "[weight_pcoa.png]", f"Result/05_Beta_diversity/PCoA/{self.groupvs}/weighted_unifrac_dmPCoA.png")

			if "[unweight_pcoa.png]" in p.text:
				self.insert_png(p, "[unweight_pcoa.png]", f"Result/05_Beta_diversity/PCoA/{self.groupvs}/unweighted_unifrac_dmPCoA.png")

			if "[nmds.png]" in p.text:
				self.insert_png(p, "[nmds.png]", f"Result/05_Beta_diversity/NMDS/{self.groupvs}/NMDS.png")

			if "[anosim.png]" in p.text:
				self.insert_png(p, "[anosim.png]", f"Result/05_Beta_diversity/anosim/{self.groupvs}.png")

			if "[cladogram.png]" in p.text:
				self.insert_png(p, "[cladogram.png]", f"Result/06_Differential_analysis/LefSe/{self.groupvs}/{self.groupvs}.cla.png")

			if "[Lda_score.png]" in p.text:
				self.insert_png(p, "[Lda_score.png]", f"Result/06_Differential_analysis/LefSe/{self.groupvs}/{self.groupvs}.png")

			if "[biomarker.png]" in p.text:
				self.insert_png(p, "[biomarker.png]", f"Result/06_Differential_analysis/LefSe/{self.groupvs}/biomarker.png")

			if "[stamp.png]" in p.text:
				self.insert_png(p, "[stamp.png]", f"Result/06_Differential_analysis/STAMP/{self.groupvs}/t-test.png")

			if "[pca12.png]" in p.text:
				self.insert_png(p, "[pca12.png]", f'Result/07_FunctionPrediction/KEGG/PCA/{self.groupvs}/pca.png',
													f'Result/07_FunctionPrediction/COG/PCA/{self.groupvs}/pca.png')

			if "[barplot12.png]" in p.text:
				self.insert_png(p, "[barplot12.png]", 'Result/07_FunctionPrediction/KEGG/Barplot/kegg_predicted_L2.png',
														'Result/07_FunctionPrediction/COG/Barplot/cog_predicted_L2.png')

			if "[heatmap12.png]" in p.text:
				self.insert_png(p, "[heatmap12.png]", 'Result/07_FunctionPrediction/KEGG/Heatmap/kegg_predicted_L2.png',
														'Result/07_FunctionPrediction/COG/Heatmap/cog_predicted_L2.png')

			if "[kegg1.png]" in p.text:
				self.insert_png(p, "[kegg1.png]", f'Result/07_FunctionPrediction/KEGG/LEfSe/{self.groupvs}/L2.png')

			if "[kegg2.png]" in p.text:
				self.insert_png(p, "[kegg2.png]", f'Result/07_FunctionPrediction/KEGG/LEfSe/{self.groupvs}/biomarker.png')

			if "[cog1.png]" in p.text:
				self.insert_png(p, "[cog1.png]", f'Result/07_FunctionPrediction/COG/LEfSe/{self.groupvs}/L2.png')

			if "[cog2.png]" in p.text:
				self.insert_png(p, "[cog2.png]", f'Result/07_FunctionPrediction/COG/LEfSe/{self.groupvs}/biomarker.png')

			if "[stamp1.png]" in p.text:
				self.insert_png(p, "[stamp1.png]", f'Result/07_FunctionPrediction/KEGG/STAMP/{self.groupvs}/t-test.png')

			if "[stamp2.png]" in p.text:
				self.insert_png(p, "[stamp2.png]", f'Result/07_FunctionPrediction/COG/STAMP/{self.groupvs}/t-test.png')


	def save(self, document):
		document.save('Report.docx')
		print(f'报告生成完成，路径为：{os.path.join(os.getcwd(), "Report.docx")}')

	def multigroups(self, index, num1, num2, path, png_type, png_list, text):
		'''
		index:html中的索引
		num1:albumSlider的索引
		num2: div的索引
		path:插入图片的路径
		png_type:插入png的类型
		png_list:插入图片的列表
		text:需要被替换掉的字符
		'''
		png_js = f'$("#{index} .albumSlider:eq({num1}) .fullview a").attr("href", "' + f'{path}{png_list[0]}' + '");'
		png_js += f'$("#{index} .albumSlider:eq({num1}) .fullview img").attr("src", "' + f'{path}{png_list[0]}' + '");'
		if len(png_list) == 1:
			png_js += f'$("#{index} .albumSlider:eq({num1}) div:eq({num2})").remove();'
		else:
			png_js += f'var {png_type} = ["' + '","'.join(png_list) + '"];'
			png_js += f'$.each({png_type}' + ',function(index,item){var $br = $("<br>");var $img = $("<img style=\'height: 60px;width: 60px\'>");'
			png_js += '$img.attr("src", "' + path + f'"+item);$("#{index} .albumSlider:eq({num1}) div:eq({num2}) center table tr").append('
			png_js += '$("<td style=\'text-align: center;\'>").append(item.replace(/' + text + '/, ""), $br, $img))});'
		return png_js

	def table_header_js(self, index, table, data, last_color):
		'''
		last_color:最后列的颜色
		'''
		table_header = data.columns.tolist()
		header_js = f'var {table} = ['
		for i in range(len(table_header)):
			if i == 0:
				header_js += '{h' + str(i) + ':"' + str(table_header[i]) + '",'
			elif all([i > 0, i < len(table_header) - 1]):
				header_js += 'h' + str(i) + ':"' + str(table_header[i]) + '",'
			else:
				header_js += 'h' + str(i) + ':"' + str(table_header[i]) + '"},'
		header_js += '];'
		header_js += '$.each(' + table + ',function(index,item){$("#' + index + ' table:eq(0)").append($("<tr>").append('

		for j in range(len(table_header)):
			if j == 0:
				header_js += f'''$('<td style="background-color: #FFC1C1">').html(item.h{str(j)}),'''
			elif all([j > 0, j < len(table_header) - 1]):
				header_js += f'''$('<td style="background-color: #B9D3EE">').html(item.h{str(j)}),'''
			else:
				header_js += f'''$('<td style="background-color: {last_color}">').html(item.h{str(j)}),'''
		header_js += '))});'

		return header_js

	def table_js(self, index, table, data):
		'''
		table:表格
		data:数据
		'''
		js_data = f'var {table} = ['
		for i in range(data.shape[0]):
			for j in range(data.shape[1]):
				if j == 0:
					js_data += '{c' + str(j) + ':"' + str(data.iloc[i, j]) + '",'
				elif all([j > 0, j < data.shape[1] - 1]):
					js_data += 'c' + str(j) + ':"' + str(data.iloc[i, j]) + '",'
				else:
					js_data += 'c' + str(j) + ':"' + str(data.iloc[i, j]) + '"},'
		js_data += '];'
		js_data += '$.each(' + table + ',function(index,item){$("#' + index + ' table:eq(0)").append($("<tr>").append('

		for j in range(data.shape[1]):
			js_data += f'$("<td>").html(item.c{str(j)}),'
		js_data += '))});'

		return js_data

	def netReport(self):
		js = os.path.join('Result', 'WebReport', 'src', 'js', 'data.js')

		venn_list = [f'{i}/venn.png' for i in os.listdir(os.path.join('Result', '02_OTU_Taxa', 'Venn'))]
		pca_list = [f'{i}/pca.png' for i in os.listdir(os.path.join('Result', '05_Beta_diversity', 'PCA'))]
		weighted_pcoa = [f'{i}/weighted_unifrac_dmPCoA.png' for i in os.listdir(os.path.join('Result', '05_Beta_diversity', 'PCoA'))]
		unweighted_pcoa = [f'{i}/unweighted_unifrac_dmPCoA.png' for i in os.listdir(os.path.join('Result', '05_Beta_diversity', 'PCoA'))]
		nmds_list = [f'{i}/NMDS.png' for i in os.listdir(os.path.join('Result', '05_Beta_diversity', 'NMDS'))]
		anosim_list = [i for i in os.listdir(os.path.join('Result', '05_Beta_diversity', 'anosim')) if re.search('png', i)]
		cla_list = [f'{i}/Cla.png' for i in os.listdir(os.path.join('Result', '06_Differential_analysis', 'LefSe'))]
		lda_list = [f'{i}/LDA.png' for i in os.listdir(os.path.join('Result', '06_Differential_analysis', 'LefSe'))]
		biomarker_list = [f'{i}/biomarker.png' for i in os.listdir(os.path.join('Result', '06_Differential_analysis', 'LefSe'))]
		stamp_list = [f'{i}/t-test.png' for i in os.listdir(os.path.join('Result', '06_Differential_analysis', 'STAMP')) if '_vs_' in i]
		kegg_pca = [f'{i}/pca.png' for i in os.listdir(os.path.join('Result', '07_FunctionPrediction', 'KEGG', 'PCA'))]
		cog_pca = [f'{i}/pca.png' for i in os.listdir(os.path.join('Result', '07_FunctionPrediction', 'COG', 'PCA'))]
		kegg_lda = [f'{i}/L2.png' for i in os.listdir(os.path.join('Result', '07_FunctionPrediction', 'KEGG', 'LefSe'))]
		kegg_biomarker = [f'{i}/biomarker.png' for i in os.listdir(os.path.join('Result', '07_FunctionPrediction', 'KEGG', 'LefSe'))]
		cog_lda = [f'{i}/L2.png' for i in os.listdir(os.path.join('Result', '07_FunctionPrediction', 'COG', 'LefSe'))]
		cog_biomarker = [f'{i}/biomarker.png' for i in os.listdir(os.path.join('Result', '07_FunctionPrediction', 'COG', 'LefSe'))]
		kegg_stamp = [f'{i}/t-test.png' for i in os.listdir(os.path.join('Result', '07_FunctionPrediction', 'KEGG', 'STAMP'))]
		cog_stamp = [f'{i}/t-test.png' for i in os.listdir(os.path.join('Result', '07_FunctionPrediction', 'COG', 'STAMP'))]

		with open(js, 'w', encoding='utf-8') as f:
			f.write('$(document).ready(function(){')
			f.write('$(".cover-info table tr:eq(0) td span").html("' + self.project_name + '");')
			f.write('$(".cover-info table tr:eq(1) td span").html("' + self.school + '");')
			f.write('$(".cover-info table tr:eq(2) td span").html("' + self.project_num + '");')
			f.write('$(".cover-info table tr:eq(3) td span").html("' + self.today + '");')

			f.write('$("#name1 .p span:eq(0)").html("' + self.total_sample + '");')
			f.write('$("#name1 .p span:eq(1)").html("' + self.group_num + '");')
			f.write('$("#name1 .p span:eq(2)").html("' + self.sample_num + '");')

			f.write(self.table_js('name1', 'sample_info', self.sample_info_df))
			f.write(self.table_js('name3', 'qc', self.qc))
			f.write(self.table_header_js('name3_2', 'otu_table_header', self.otu_table, '#FFC1C1'))
			f.write(self.table_js('name3_2', 'otu_table', self.otu_table))
			f.write(self.multigroups('name3_2_2', 0, 1,  '../../02_OTU_Taxa/Venn/', 'venn', venn_list, '\/venn.png'))
			f.write('$("#name3_3_2 .albumSlider .fullview a").attr("href", "' + f'../../03_Community/taxa_tree/sample_tree/{self.sample_tree}' + '");')
			f.write('$("#name3_3_2 .albumSlider .fullview img").attr("src", "' + f'../../03_Community/taxa_tree/sample_tree/{self.sample_tree}' + '");')
			f.write(self.table_js('name3_4', 'alpha', self.alpha_index))
			f.write(self.multigroups('name3_5_2', 0, 1, '../../05_Beta_diversity/PCA/', 'pca', pca_list, 'pca.png'))
			f.write(self.multigroups('name3_5_3', 0, 1, '../../05_Beta_diversity/PCoA/', 'weighted_pcoa', weighted_pcoa, '\/weighted_unifrac_dmPCoA.png'))
			f.write(self.multigroups('name3_5_3', 1, 1, '../../05_Beta_diversity/PCoA/', 'unweighted_pcoa', unweighted_pcoa, '\/unweighted_unifrac_dmPCoA.png'))
			f.write(self.multigroups('name3_5_4', 0, 1, '../../05_Beta_diversity/NMDS/', 'nmds', nmds_list, '\/NMDS.png'))
			f.write(self.table_js('name3_6', 'anosim', self.anosim))
			f.write(self.multigroups('name3_6', 0, 1, '../../05_Beta_diversity/anosim/', 'anosim', anosim_list, '.png'))
			f.write(self.table_js('name3_6_2', 'adonis', self.adonis))
			f.write(self.multigroups('name3_7', 0, 1, '../../06_Differential_analysis/LefSe/', 'cla', cla_list, '\/Cla.png'))
			f.write(self.multigroups('name3_7', 1, 1, '../../06_Differential_analysis/LefSe/', 'lda', lda_list, '\/LDA.png'))
			f.write(self.multigroups('name3_7', 2, 1, '../../06_Differential_analysis/LefSe/', 'biomarker', biomarker_list, '\/biomarker.png'))
			f.write(self.multigroups('name3_7_2', 0, 1, '../../06_Differential_analysis/STAMP/', 'stamp', stamp_list, '\/t-test.png'))
			f.write(self.table_header_js('name3_8', 'ko_header', self.ko, '#B9D3EE'))
			f.write(self.table_js('name3_8', 'ko', self.ko))
			f.write(self.table_header_js('name3_8_2', 'cog_header', self.cog, '#B9D3EE'))
			f.write(self.table_js('name3_8_2', 'cog', self.cog))
			f.write(self.multigroups('name3_8_3', 0, 1, '../../07_FunctionPrediction/KEGG/PCA/', 'kegg_pca', kegg_pca, '\/pca.png'))
			f.write(self.multigroups('name3_8_3', 1, 1, '../../07_FunctionPrediction/COG/PCA/', 'cog_pca', cog_pca, '\/pca.png'))
			f.write(self.multigroups('name3_8_6', 0, 1, '../../07_FunctionPrediction/KEGG/LefSe/', 'kegg_lda', kegg_lda, '\/L2.png'))
			f.write(self.multigroups('name3_8_6', 1, 1, '../../07_FunctionPrediction/KEGG/LefSe/', 'kegg_biomarker', kegg_biomarker, '\/biomarker.png'))
			f.write(self.multigroups('name3_8_6', 2, 1, '../../07_FunctionPrediction/COG/LefSe/', 'cog_lda', cog_lda, '\/L2.png'))
			f.write(self.multigroups('name3_8_6', 3, 1, '../../07_FunctionPrediction/COG/LefSe/', 'cog_biomarker', cog_biomarker, '\/biomarker.png'))
			f.write(self.multigroups('name3_8_7', 0, 1, '../../07_FunctionPrediction/KEGG/STAMP/', 'kegg_stamp', kegg_stamp, '\/t-test.png'))
			f.write(self.multigroups('name3_8_7', 1, 1, '../../07_FunctionPrediction/COG/STAMP/', 'cog_stamp', cog_stamp, '\/t-test.png'))

			f.write('})')

