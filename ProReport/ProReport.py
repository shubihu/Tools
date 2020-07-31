# -*- coding: future_fstrings -*-     # should work even without -*-

import os
import sys
import click
from docx import Document
from templatePy.update import get_version
from templatePy.update import update

version = 1.16
newversion = get_version()

if version < newversion:
	print('服务器上发现最新版本，开始下载更新代码。\n服务器(192.168.130.252)上代码地址：/database/proteome/ProReport')
	update(os.path.abspath(os.path.dirname(sys.argv[0])))
	exit()

platform = sys.platform

def labfree(*args):
	input_path, types, document, paragraphs, tables, nobioinfo = args
	from templatePy.project import Labfree
	labfree = Labfree(input_path, types)
	if nobioinfo:
		labfree.nobioinfo(paragraphs)
	labfree.header(paragraphs)
	labfree.table_data(tables)
	labfree.text_png_data(document, paragraphs)
	labfree.save(document)
	if platform == 'win32':
		if nobioinfo:
			labfree.update()

def itrq(*args):
	input_path, types, document, paragraphs, tables, nobioinfo = args
	from templatePy.project import Itraq
	itraq = Itraq(input_path, types)
	if nobioinfo:
		itraq.nobioinfo(paragraphs)
	itraq.header(paragraphs, start_row=13)
	itraq.table_data(tables)
	itraq.text_png_data(document, paragraphs)
	itraq.save(document)
	if platform == 'win32':
		if nobioinfo:
			itraq.update()

def tmt(*args):
	input_path, types, document, paragraphs, tables, nobioinfo = args
	from templatePy.project import TMT
	tmt = TMT(input_path, types)
	if nobioinfo:
		tmt.nobioinfo(paragraphs)
	tmt.header(paragraphs, start_row=13)
	tmt.table_data(tables)
	tmt.text_png_data(document, paragraphs)
	tmt.save(document)
	if platform == 'win32':
		if nobioinfo:
			tmt.update()

def dia(*args):
	input_path, _, document, paragraphs, tables, _ = args
	from templatePy.dia import DIA
	dia = DIA(input_path)
	dia.header(paragraphs)
	dia.table_data(tables)
	dia.text_png_data(paragraphs)
	dia.save(document)

def pholabfree(*args):
	input_path, types, document, paragraphs, tables, _ = args
	from templatePy.project import PhoLabfree
	pholabfree = PhoLabfree(input_path, types)
	pholabfree.header(paragraphs)
	pholabfree.table_data(tables)
	pholabfree.text_png_data(document, paragraphs)
	pholabfree.save(document)

def glylabfree(*args):
	input_path, types, document, paragraphs, tables, _ = args
	from templatePy.project import GlyLabfree
	glylabfree = GlyLabfree(input_path, types)
	glylabfree.header(paragraphs, start_row=5)
	glylabfree.table_data(tables)
	glylabfree.text_png_data(document, paragraphs)
	glylabfree.save(document)

def nlabfree(*args):
	input_path, types, document, paragraphs, tables, _ = args
	from templatePy.project import NLabfree
	nlabfree = NLabfree(input_path, types)
	nlabfree.header(paragraphs, start_row=5)
	nlabfree.table_data(tables)
	nlabfree.text_png_data(document, paragraphs)
	nlabfree.save(document)

def suclabfree(*args):
	input_path, types, document, paragraphs, tables, _ = args
	from templatePy.project import SucLabfree
	suclabfree = SucLabfree(input_path, types)
	suclabfree.header(paragraphs, start_row=5)
	suclabfree.table_data(tables)
	suclabfree.text_png_data(document, paragraphs)
	suclabfree.save(document)

def tyrlabfree(*args):
	input_path, types, document, paragraphs, tables, _ = args
	from templatePy.project import TyrLabfree
	tyrlabfree = TyrLabfree(input_path, types)
	tyrlabfree.header(paragraphs, start_row=5)
	tyrlabfree.table_data(tables)
	tyrlabfree.text_png_data(document, paragraphs)
	tyrlabfree.save(document)

def acelabfree(*args):
	input_path, types, document, paragraphs, tables, _ = args
	from templatePy.project import AceLabfree
	acelabfree = AceLabfree(input_path, types)
	acelabfree.header(paragraphs, start_row=5)
	acelabfree.table_data(tables)
	acelabfree.text_png_data(document, paragraphs)
	acelabfree.save(document)

def mallabfree(*args):
	input_path, types, document, paragraphs, tables, _ = args
	from templatePy.project import MalLabfree
	mallabfree = MalLabfree(input_path, types)
	mallabfree.header(paragraphs, start_row=5)
	mallabfree.table_data(tables)
	mallabfree.text_png_data(document, paragraphs)
	mallabfree.save(document)

def photmt(*args):
	input_path, types, document, paragraphs, tables, _ = args
	from templatePy.project import PhoTMT
	photmt = PhoTMT(input_path, types)
	photmt.header(paragraphs, start_row=19)
	photmt.table_data(tables)
	photmt.text_png_data(document, paragraphs)
	photmt.save(document)

def acetmt(*args):
	input_path, types, document, paragraphs, tables, _ = args
	from templatePy.project import AceTMT
	acetmt = AceTMT(input_path, types)
	acetmt.header(paragraphs, start_row=19)
	acetmt.table_data(tables)
	acetmt.text_png_data(document, paragraphs)
	acetmt.save(document)

def ipi(*args):
	input_path, _, document, paragraphs, tables, _ = args
	from templatePy.conjoint import Conjoint
	ipi = Conjoint(input_path)
	ipi.header(paragraphs)
	ipi.table_data(tables)
	ipi.text_png_data(paragraphs)
	ipi.save(document)


@click.command()
@click.option("--input_path", '-i', required=True, help="项目路径")
@click.option("--types", '-t', required=True, help='''项目类型, l：labfree; i：Itraq; t：TMT; d：DIA; pl：磷酸化Labfree; gl:泛素化labfree;
				nl:糖基化labfree; sl:琥珀酰化labfree; yl:络氨酸磷酸化labfree; al:乙酰化labfree; ml:丙二酰化labfree; pt：磷酸化TMT; 
				at：乙酰化TMT; ipi:蛋白磷酸化联合''')
@click.option('--nobioinfo/--no-nobioinfo', '-y/-n', default=False, help='是否出无生信版报告，默认为否')
def main(input_path, types, nobioinfo):
	if nobioinfo:
		print('无生信版报告')
	types = types.lower()
	types_dict = {'l': 'Labelfree报告模板-20200309.docx',
				  'i': 'iTRAQ报告模板-20200325-4标8标.docx',
				  't': 'TMT报告模板-20200325-6标10标16标.docx',
				  'd': 'DIA蛋白质组研究正式实验报告模板.docx',
				  'pl': '磷酸化Labelfree报告模板无激酶分析版.docx',
				  'gl': '泛素化label free-报告模板调整-20200517.docx',
				  'nl': 'N-糖基化label free-报告模板调整-20200520.docx',
				  'sl': '琥珀酰化label free-报告模板调整-20200517.docx',
				  'al': '乙酰化label free-报告模板调整-20200517.docx',
				  'ml': '丙二酰化label free-报告模板调整-20200517.docx',
				  'yl': '酪氨酸磷酸化label free-报告模板调整-20200517.docx',
				  'pt': '磷酸化TMT6plex 10plex TMT16相对定量蛋白质组学报告模板-自动化用20200519.docx',
				  'at': '乙酰化TMT6plex 10plex TMT16相对定量蛋白质组学报告模板-自动化用20200519.docx',
				  'ipi': '蛋白+磷酸化联合分析项目报告模板-定稿20200513.docx'}

	project_dict = {'l': labfree, 'i': itrq, 't': tmt, 'd': dia, 'pl': pholabfree, 'gl': glylabfree, 'nl': nlabfree, 'sl': suclabfree,
					'yl': tyrlabfree, 'al': acelabfree,'ml': mallabfree, 'pt': photmt, 'at': acetmt, 'ipi': ipi}

	template_file = os.path.join(os.path.abspath(os.path.dirname(sys.argv[0])), 'template', types_dict.get(types))
	document = Document(template_file)
	paragraphs = document.paragraphs
	tables = document.tables

	project_dict.get(types)(input_path, types, document, paragraphs, tables, nobioinfo)

if __name__ == '__main__':
	main()
	

