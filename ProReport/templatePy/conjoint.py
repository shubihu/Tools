# -*- coding: future_fstrings -*-     # should work even without -*-
import os
import re
import sys
import math
import datetime
import pandas as pd
pd.options.mode.chained_assignment = None
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from templatePy.template import extract_top
from templatePy.template import Template

class Conjoint(Template):
	def __init__(self, path):
		os.chdir(path)

		information_file = [i for i in os.listdir('.') if re.search('infor?mation', i, re.I)]
		if information_file:
			projectinfo = pd.read_excel(information_file[0], header=None)
		else:
			print('Error:该项目下无project information表')
			sys.exit()
		self.school = projectinfo.iloc[0, 1]
		self.project_name = projectinfo.iloc[1, 1]
		self.project_num = projectinfo.iloc[2, 1]
		self.groupvs = projectinfo.iloc[3, 1]

		pro_count = pd.read_excel(os.path.join('Statistics', '表1-关联蛋白数量统计表.xlsx'))
		self.pro_count1 = pro_count.iloc[:, :5]
		self.pro_count1.loc[''] = ''
		self.pro_count2 = pro_count.iloc[:, 5:]
		self.pro_count2.insert(0,'groupvs', pro_count.iloc[:, 0])

		pro_diff = pd.read_excel(os.path.join('Statistics', '表2-关联蛋白差异表达统计表.xlsx'), index_col=0, header=None)
		pro_diff.columns = pro_diff.loc['Comparisons']
		self.pro_diff = pro_diff[self.groupvs]
		
		pro_detail = pd.read_excel(os.path.join(self.groupvs, '表3-关联蛋白差异表达详情表.xlsx')).head(5)
		self.pro_detail = pro_detail.iloc[:, :9]
		self.pro_detail = self.pro_detail.applymap(lambda x: round(x, 3) if isinstance(x, float) else x)
		self.pro_detail['Phosphorylated Site'] = self.pro_detail['Phosphorylated Site'].astype(int)

		self.ppi = pd.read_excel(os.path.join(self.groupvs, 'PPI', '表6-PPI互作分析表.xlsx')).head(5)

		kinase = pd.read_excel(os.path.join(self.groupvs, 'Kinase', '表7-关联激酶数量统计表.xlsx'))
		self.kinase1 = kinase.iloc[:, :4]
		self.kinase1.loc[''] = ''
		self.kinase2 = kinase.iloc[:, 4:]

		self.kinase_substrate = pd.read_excel(os.path.join(self.groupvs, 'Kinase', '表8-激酶底物对应关系详情表.xlsx')).head(5)
		self.kinase_substrate = self.kinase_substrate.applymap(lambda x: round(x, 3) if isinstance(x, float) else x)

		kegg = pd.read_excel(os.path.join(self.groupvs, 'bothPMP', 'kegg', 'kegg.xlsx'), sheet_name=None)
		self.keggID = extract_top(kegg, 'keggID')[0][0]


	def header(self, paragraphs):
		today = str(datetime.date.today())
		pa = paragraphs[16].add_run(self.school)
		self.paragraph_format(pa, size=14, family=u'微软雅黑')
		pa = paragraphs[17].add_run(self.project_name)
		self.paragraph_format(pa, size=14, family=u'微软雅黑')
		pa = paragraphs[18].add_run(self.project_num)
		self.paragraph_format(pa, size=14, family=u'微软雅黑')
		pa = paragraphs[19].add_run(today)
		self.paragraph_format(pa, size=14, family="微软雅黑")

	def table_data(self, tables):
		self.insert_table(self.pro_count1, tables[0], size=10, family_ch=u'微软雅黑', family_en='微软雅黑')
		df_tb1 = pd.DataFrame([['', 'Correlated proteins', 'Correlated proteins(Differential expressed only in protein level)',
								'Correlated proteins(Differential expressed only in phosphoproteins level)',
								'Correlated proteins(Differential expressed both in protein and phosphoprotein level)'], [''] * 5])
		self.insert_table(df_tb1, tables[0], size=10, family_ch=u'微软雅黑', family_en='微软雅黑', bold=True)
		self.insert_table(self.pro_count2, tables[0], size=10, family_ch=u'微软雅黑', family_en='微软雅黑')

		tmp_text = tables[1].cell(0, 1).text
		tmp_text = tmp_text.replace('groupvs', str(self.groupvs))
		tables[1].cell(0, 1).text = ''
		tables[1].cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		tables[1].cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		self.paragraph_format(tables[1].cell(0, 1).paragraphs[0].add_run(tmp_text), size = 9, family = 'Arial', bold=True)

		for i in range(3, 6):
			for j in range(1, 4):
				tables[1].cell(i, j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
				tables[1].cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
				self.paragraph_format(tables[1].cell(i, j).paragraphs[0].add_run(str(self.pro_diff.iloc[i - 1, j - 1])), size = 9, family = 'Arial')

		self.insert_table(self.pro_detail, tables[2], size=7, family_ch=u'微软雅黑', family_en='微软雅黑')
		self.insert_table(self.ppi, tables[3], size=10, family_ch=u'微软雅黑', family_en='微软雅黑')

		df_tb2 = pd.DataFrame([['Correlated kinases', 'Correlated kinases (Differential expressed only in protein level)',
								'Correlated kinases (Differential expressed only in phosphorylated kinases level)',
								'Correlated proteins(Differential expressed both in kinases and phosphorylated kinases level)'], [''] * 4])
		self.insert_table(self.kinase1, tables[4], size=10, family_ch=u'微软雅黑', family_en='微软雅黑')
		self.insert_table(df_tb2, tables[4], size=10, family_ch=u'微软雅黑', family_en='微软雅黑', bold=True)
		self.insert_table(self.kinase2, tables[4], size=10, family_ch=u'微软雅黑', family_en='微软雅黑')

		self.insert_table(self.kinase_substrate, tables[5], size=7, family_ch=u'微软雅黑', family_en='微软雅黑')

	def text_png_data(self, paragraphs):

		for i, p in enumerate(paragraphs):
			
			if "[venn]" in p.text:
				replace = self.insert_png(p, "[venn]", os.path.join(self.groupvs, 'Venn.png'))
				if replace == False:
					self.text_replace(p, ["[venn]"],['无venn图'], family_ch=u'微软雅黑', family_en='微软雅黑')
			if '[volcano]' in p.text:
				self.insert_png(p, '[volcano]', os.path.join(self.groupvs, 'Volcano_Protein2.png'))
			if "[cluster]" in p.text:
				self.insert_png(p, "[cluster]", os.path.join(self.groupvs, 'FCcluster.png'))
			if "[go_enrich1]" in p.text:
				self.insert_png(p, "[go_enrich1]", os.path.join(self.groupvs, 'GO', 'GO_Enrichment.png'))
			if '[go_level2-1]' in p.text:
				self.insert_png(p, '[go_level2-1]', os.path.join(self.groupvs, 'onlyMP', 'go', 'GOLevel2.png'))
			if '[go_level2-2]' in p.text:
				self.insert_png(p, '[go_level2-2]', os.path.join(self.groupvs, 'bothPMP', 'go', 'GOLevel2.png'))
			if "[go_enrich2]" in p.text:
				self.insert_png(p, "[go_enrich2]", os.path.join(self.groupvs, 'onlyMP', 'go', 'GO_Enrichment.png'))
			if "[go_enrich3]" in p.text:
				self.insert_png(p, "[go_enrich3]", os.path.join(self.groupvs, 'bothPMP', 'go', 'GO_Enrichment.png'))
			if "[kegg_enrich1]" in p.text:
				self.insert_png(p, "[kegg_enrich1]", os.path.join(self.groupvs, 'KEGG', 'KEGG_Enrichment.png'))
			if '[top_map1]' in p.text:
				self.insert_png(p, '[top_map1]', os.path.join(self.groupvs, 'onlyMP', 'kegg', 'TopMapStat.png'))
			if '[top_map2]' in p.text:
				self.insert_png(p, '[top_map2]', os.path.join(self.groupvs, 'bothPMP', 'kegg', 'TopMapStat.png'))
			if "[kegg_enrich2]" in p.text:
				self.insert_png(p, "[kegg_enrich2]", os.path.join(self.groupvs, 'onlyMP', 'kegg', 'KEGG_Enrichment.png'))
			if "[kegg_enrich3]" in p.text:
				self.insert_png(p, "[kegg_enrich3]", os.path.join(self.groupvs, 'bothPMP', 'kegg', 'KEGG_Enrichment.png'))
			if "[pathway]" in p.text:
				self.insert_png(p, "[pathway]", os.path.join(self.groupvs, 'bothPMP', 'kegg', 'map', f'{self.keggID}.png'))
			if "[ppi]" in p.text:
				self.insert_png(p, "[ppi]", os.path.join(self.groupvs, 'PPI', 'ppi.png'))
			if '[venn_kinase]' in p.text:
				self.insert_png(p, '[venn_kinase]', os.path.join(self.groupvs, 'Kinase', 'Venn_Kinase.png'))
			if '[volcano_kinase]' in p.text:
				self.insert_png(p, '[volcano_kinase]', os.path.join(self.groupvs, 'Kinase', 'Volcano_Kinase.png'))
			if '[cluster_kinase]' in p.text:
				self.insert_png(p, '[cluster_kinase]', os.path.join(self.groupvs, 'Kinase', 'FCcluster_Kinase.png'))
			if '[chord]' in p.text:
				self.insert_png(p, '[chord]', os.path.join(self.groupvs, 'Kinase', 'Chord.png'))

	def save(self, document):
		document.save('蛋白磷酸化联合分析报告.docx')
		print('报告生成完成')





