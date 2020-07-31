# -*- coding: future_fstrings -*-     # should work even without -*-

import os
import sys
import re
import datetime
import pandas as pd
pd.options.mode.chained_assignment = None
# import modin.pandas as pd
from collections import Counter
from collections import OrderedDict
# from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
# import win32com.client

def extract_top(data, func, num=5):
		top = []
		pathway = []
		if func == 'BP':
			data = data['Enrichment']
			data = data[data['Category'] == 'P'].head(num)
			top = data['Term'].tolist()
		elif func == 'MF':
			data = data['Enrichment']
			data = data[data['Category'] == 'F'].head(num)
			top = data['Term'].tolist()
		elif func == 'CC':
			data = data['Enrichment']
			data = data[data['Category'] == 'C'].head(num)
			top = data['Term'].tolist()
		elif func == 'goEnrich':
			data = data['Enrichment'].head(num)
			top = data['Term'].tolist()
		elif func == 'keggEnrich':
			data = data['Enrichment'].head(num)
			top = data['Map_Name'].tolist()
			pathway = top[0]
		elif func == 'keggID':
			data = data['Enrichment'].head(num)
			top = data['Map_ID'].tolist()
		elif func == 'map2query':
			data = data['map2query'].head(num)
			top = data['Map_Name'].tolist()

		if not func in ['keggID', 'map2query']:
			pvalue_list = data['P value'].tolist()
			pvalue_top = [i for i in pvalue_list if float(i) < 0.05]
			top = top[:len(pvalue_top)] if pvalue_top else []
		return top, pathway


class Template:
	def __init__(self, path, types):
		self.types = types
		os.chdir(path)

		information_file = [i for i in os.listdir('.') if re.search('infor?mation', i, re.I)]
		if information_file:
			if 'xls' in information_file[0]:
				self.projectinfo = pd.read_excel(information_file[0], index_col=0, header=None)
			elif information_file[0].endswith('csv'):
				self.projectinfo = pd.read_csv(information_file[0], index_col=0, header=None)
			elif information_file[0].endswith('txt'):
				self.projectinfo = pd.read_csv(information_file[0], index_col=0, sep='\t', header=None)
		else:
			raise Exception('Error:该项目下无project_information表')

		# index_list = ['项目名称', '委托单位', '项目编号', '物种', '比较组', '数据库', '样品组数', '每组生物学重复数', ' 样品总数']
		# self.projectinfo.index = index_list
		index_list = self.projectinfo.index.tolist()
		groupvs = [str(i) for i in index_list if re.search('比较组', str(i))]
		# index_list[index_list.index(groupvs[0])] = '比较组'
		# self.projectinfo.index = index_list

		self.species = self.projectinfo.loc['物种'][1]
		self.groupvs = str(self.projectinfo.loc[groupvs[0]][1])
		self.database = self.projectinfo.loc['数据库'][1]
		
		self.sampleInfo = pd.read_csv('samples.txt', sep='\t', header=None)
		self.origi_record = pd.read_excel('原始记录.xlsx', sheet_name=1).fillna('')
		self.fc = float(self.origi_record[self.origi_record.iloc[:, 0] == '差异倍数'].iloc[0, 1])
		if self.types in ['pl', 'gl', 'nl', 'sl', 'yl', 'al', 'ml', 'pt', 'at']:
			statistic_file = [i for i in os.listdir('Evaluation') if re.search('^(.(?!_))*Statistic.*csv$', i, re.I)]
			if statistic_file:
				statistic = pd.read_csv(os.path.join('Evaluation', statistic_file[0]))
		else:
			statistic = pd.read_csv(os.path.join('Evaluation', 'Statistic.csv'))
		self.statistic_list = [str(j) for i in statistic.values.T.tolist() for j in i]

		peptideScore_file = [i for i in os.listdir('Evaluation') if re.search('PeptideScore', i, re.I)]
		if peptideScore_file:
			peptideScore = os.path.join('Evaluation', peptideScore_file[0])
			with open(peptideScore) as f:
				self.medianScore = f.readline().split('=')[-1].strip()
				self.percentage = re.split('[=(]', f.readline())[1].strip()

		if self.groupvs != 'nan':
			go_file = os.path.join(self.groupvs, 'go', 'GO.xlsx')
			self.go = pd.read_excel(go_file, sheet_name=None)
			self.goEnrich_top5 = extract_top(self.go, 'goEnrich')[0]

			kegg_file = os.path.join(self.groupvs, 'kegg', 'kegg.xlsx')
			self.kegg = pd.read_excel(kegg_file, sheet_name=None)
			self.keggEnrich_top5, self.pathway = extract_top(self.kegg, 'keggEnrich')
		else:
			self.go = self.goEnrich_top5 = None
			self.kegg = self.keggEnrich_top5 = self.pathway = None

	
	def paragraph_format(self, pa, size, family, r = 0x00, g = 0x00, b = 0x00, bold = None):
		pa.font.size = Pt(size)
		pa.font.name = family
		if bold == True:
			pa.font.bold = True
		pa.font.color.rgb = RGBColor(r, g, b)
		pa._element.rPr.rFonts.set(qn('w:eastAsia'), family)

	def delete_paragraph(self, paragraphs, p_index_list):
		'''
		p_index_list:需要删除段落索引的列表
		'''
		for i in p_index_list:
			p = paragraphs[i]
			p = p._element
			if p.getparent() is not None:
				p.getparent().remove(p)
				p._p = p._element = None

	def text_replace(self, p, text1_list, text2_list, size=10.5, family_ch=u'微软雅黑', family_en='Arial', bold=None):
		'''
		p:段落
		text1_list: 需要替换的文本列表
		text2_list: 替换的文本列表,与上面列表长度相等
		'''
		text1_list = [i.strip('[]') for i in text1_list]
		text1_list = text1_list[0] if len(text1_list) == 1 else '|'.join(text1_list)
		text = p.text.strip()
		text_list = re.split(text1_list, text)
		p.clear()
		for i in range(len(text_list) - 1):
			self.paragraph_format(p.add_run(text_list[i]), size = size, family = family_ch, bold=bold)
			if '无P值小于0.05' in text2_list[i]:
				self.paragraph_format(p.add_run(text2_list[i]), size = size, family = family_ch, bold=bold)
			else:
				self.paragraph_format(p.add_run(text2_list[i]), size = size, family = family_en, bold=bold)
		self.paragraph_format(p.add_run(text_list[-1]), size = size, family = family_ch, bold=bold)

	def insert_table(self, data, table, axis=0, size=9, family_ch=u'微软雅黑', family_en='Arial', bold=None):
		'''
		data:插入的数据，格式为dataframe
		table:待插入数据的表格
		axis : 0为行添加，1为列添加
		'''
		total_row = len(table.rows)
		total_columns = len(table.columns)
		if axis == 0:
			row_cells = table.rows[total_row - 1].cells
			for row_num in range(data.shape[0]):
				if row_num > 0:
					row_line = table.add_row()
					row_line.height = Cm(0.7)
					row_cells = row_line.cells
					# row_cells = table.rows[total_row + row_num -1].cells			
				for col_num in range(data.shape[1]):
					tmp = data.iloc[row_num, col_num]
					if type(tmp) == 'float':
						tmp = int(tmp)
					pa = row_cells[col_num].paragraphs[0].add_run(str(tmp))
					row_cells[col_num].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					row_cells[col_num].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=size, family=family_en, bold=bold)
		else:
			if data.shape[0] > 2:
				for i in range(data.shape[0] - 2):
					row_line = table.add_row()
					row_line.height = Cm(1)
					row_cells = row_line.cells
					pa = row_cells[0].paragraphs[0].add_run('样本名称')
					row_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=size, family=family_ch, bold=bold)

			col_cells = table.columns[total_columns - 1].cells
			for col_num in range(data.shape[1]):
				if col_num > 0:
					col_line = table.add_column(Inches(0.7))
					col_cells = col_line.cells
					col_cells = table.columns[total_columns + col_num -1].cells
				for row_num in range(data.shape[0]):
					tmp = data.iloc[row_num, col_num]
					pa = col_cells[row_num].paragraphs[0].add_run(str(tmp))
					col_cells[row_num].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					col_cells[row_num].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					self.paragraph_format(pa, size=size, family=family_en, bold=bold)

	def table_center(self, table):
		'''
		table 表格数据居中
		'''
		for i in range(len(table.rows)):
			for j in range(len(table.columns)):
				table.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
				table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# return table

	def move_table_after(self, table, paragraph):
		'''
		移动表格到指定位置
		'''
		tbl, p = table._tbl, paragraph._p
		p.addnext(tbl)

	def insert_png(self, p, png1, png2, png3=None):
		if os.path.exists(png2):
			p.text = p.text.strip().replace(png1, '')
			run = p.add_run()
			if png3:
				run.add_picture(png2, width=Inches(2.5))
				run.add_picture(png3, width=Inches(2.5))
			else:
				run.add_picture(png2, width=Inches(4.5))
			return True
		else:
			return False

	def record(self, record):		
		record_diff = record[record.loc[:, 'group'].str.contains('vs|oneway|twoway')]
		# f1 = lambda x: re.sub(r'\(.*\)', '', str(x))
		# f2 = lambda x: int(float(x))
		tmp = record_diff.iloc[:, 1:].applymap(lambda x: re.sub(r'\(.*\)', '', str(x)))  # applymap对每个元素进行处理
		tmp = tmp.applymap(lambda x: x if '-' in str(x) else int(float(x)))
		record_diff = pd.concat([record_diff.iloc[:, 0], tmp], axis=1)
		return record_diff

	def header(self, paragraphs, start_row=10):
		today = str(datetime.date.today())
		pa = paragraphs[start_row].add_run(self.projectinfo.loc['项目名称'][1])
		self.paragraph_format(pa, size=14, family=u'微软雅黑', bold=True)
		pa = paragraphs[start_row + 1].add_run(self.projectinfo.loc['委托单位'][1])
		self.paragraph_format(pa, size=14, family=u'微软雅黑', bold=True)
		pa = paragraphs[start_row + 2].add_run(self.projectinfo.loc['项目编号'][1])
		self.paragraph_format(pa, size=14, family="Arial", bold=True) #### family = 'Calibri'
		pa = paragraphs[start_row + 3].add_run(today)
		self.paragraph_format(pa, size=14, family="Arial", bold=True)

	# =============================================================================
	# 插入表格数据
	# =============================================================================

	def table_data(self, tables):
		summary1 = tables[0]
		self.paragraph_format(summary1.cell(2,0).paragraphs[0].add_run(self.species), size = 9, family = 'Arial')
		self.paragraph_format(summary1.cell(2,1).paragraphs[0].add_run(self.database), size = 9,family = 'Arial')
		
		# sampleInfo.iloc[:, 0] = sampleInfo.iloc[:, 0].apply(lambda x: re.sub('-\d+$', '', x))
		def ab(df):return';'.join(df.values)
		samples_data = pd.DataFrame(self.sampleInfo.groupby(2)[1].apply(ab))  ## 多行合并一行
		samples_data.insert(0,'', samples_data.index.tolist())
		self.insert_table(samples_data, tables[1])
		
		statistic_df = pd.DataFrame(self.statistic_list).T
		self.insert_table(statistic_df, tables[2])
		self.insert_table(statistic_df, tables[5])

		diff_table = tables[3]
		tmp_text = diff_table.cell(1, 1).text
		if 'upRatio' in tmp_text:
			tmp_text = tmp_text.replace('upRatio', str(self.fc))
			diff_table.cell(1, 1).text = ''
			diff_table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
			diff_table.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			self.paragraph_format(diff_table.cell(1, 1).paragraphs[0].add_run(tmp_text), size = 9, family = 'Arial', bold=True)
		record_diff = self.record(self.origi_record)
		self.insert_table(record_diff, tables[3])

		func_table = tables[4]
		self.paragraph_format(func_table.cell(2,0).paragraphs[0].add_run(self.groupvs), size = 9, family = 'Arial')
		self.paragraph_format(func_table.cell(8,0).paragraphs[0].add_run(self.groupvs), size = 9, family = 'Arial')
		if self.goEnrich_top5:
			for i in range(len(self.goEnrich_top5)):
				self.paragraph_format(func_table.cell(i+2,1).paragraphs[0].add_run(self.goEnrich_top5[i]),size = 9, family ='Arial')
		if self.keggEnrich_top5:
			for i in range(len(self.keggEnrich_top5)):
				self.paragraph_format(func_table.cell(i+8,1).paragraphs[0].add_run(self.keggEnrich_top5[i]),size = 9, family ='Arial')
				
		if self.types in ['l']:
			protein = pd.read_excel(os.path.join('质谱鉴定和定量结果', '附件1_蛋白质鉴定列表.xlsx'), sheet_name=0)
			LFQ_list = [i for i in protein.columns if re.search('LFQ', i)]
			sample = [i.split(' ')[-1] for i in LFQ_list]
			sample.insert(0, 'Total')
			proNum = [protein[i].count() for i in LFQ_list]
			proNum.insert(0, protein['Protein'].count())
			database_list = [self.database] * len(sample)
			data_frame = pd.DataFrame(columns=['database', 'sample', 'proNum'])
			data_frame['database'] = database_list
			data_frame['sample'] = sample
			data_frame['proNum'] = proNum
			self.insert_table(data_frame, tables[6])
			self.insert_table(record_diff, tables[7])
			pa = tables[8].cell(8,1).paragraphs[0].add_run(self.database)
			self.paragraph_format(pa, size=9, family="Arial")
		if self.types in ['i', 't', 'pt', 'at']:
			self.insert_table(record_diff, tables[6])
			pa = tables[8].cell(7,1).paragraphs[0].add_run(self.database)
			self.paragraph_format(pa, size=9, family="Arial")

			itraq_tmt = self.sampleInfo.T.iloc[:2,:]
			itraq_tmt.iloc[0, :] = itraq_tmt.iloc[0, :].apply(lambda x: re.split('[. ]', str(x))[-1])
			if len(itraq_tmt.iloc[0,:]) != len(set(itraq_tmt.iloc[0, :])):
				newdf = pd.DataFrame()
				for i, j in itraq_tmt.T.groupby(0):
					newdf[i] = j[1].tolist()
				newdf = newdf.T
				newdf.insert(0, '', newdf.index.tolist())
				newdf = newdf.T
				newdf.reset_index(inplace=True)
				newdf.drop('index', axis =1, inplace=True)
				No = newdf.index.tolist()
				No[0] = 'No.'
				newdf['No'] = No
				self.insert_table(newdf, tables[7], axis=1)
			else:
				No = itraq_tmt.index.tolist()
				No[0] = 'No.'
				itraq_tmt['No'] = No
				self.insert_table(itraq_tmt, tables[7], axis=1)
		if self.types in ['pl', 'gl', 'nl', 'sl', 'yl', 'al', 'ml']:
			self.insert_table(record_diff, tables[6])
			pa = tables[7].cell(8,1).paragraphs[0].add_run(self.database)
			self.paragraph_format(pa, size=9, family="Arial")

	# =============================================================================
	# 插入文本及结果图片
	# =============================================================================
	def text_png_data(self, document, paragraphs):
		if os.path.exists(os.path.join('Evaluation', 'Venn')):
			# if self.types in ['pl', 'gl']:
			# 	venn_pro = [i for i in os.listdir(os.path.join('Evaluation', 'Venn', '组间', 'ProteinVenn')) if re.search('png', i)]
			# 	venn_pep = [i for i in os.listdir(os.path.join('Evaluation', 'Venn', '组间', 'PeptideVenn')) if re.search('png', i)]
			# else:
			venn2 = [i for i in os.listdir(os.path.join('Evaluation', 'Venn', '组间')) if re.search('png', i)]
		if all([self.go, self.kegg]):
			keggID = extract_top(self.kegg, 'keggID')[0][0]
			bp_top = extract_top(self.go, 'BP')[0]
			mf_top = extract_top(self.go, 'MF')[0]
			cc_top = extract_top(self.go, 'CC')[0]
			map2query_top5 = extract_top(self.kegg, 'map2query')[0]
		else:
			keggID = bp_top = mf_top = cc_top = map2query_top5 = ''
		groupNum = str(len(set(self.sampleInfo.iloc[:,2])))
		Num = list(Counter(self.sampleInfo.iloc[:,2]).values())
		Num = [str(i) for i in Num]
		Num = ','.join(list(set(Num))) if len(set(Num)) > 1 else str(Num[0])
		total_num = str(len(self.sampleInfo.iloc[:,2]))
		if os.path.exists(os.path.join('Evaluation', 'ModifiedSiteAnnot.txt')):
			with open(os.path.join('Evaluation', 'ModifiedSiteAnnot.txt')) as fs:
					tmp1 = fs.readline().strip().split('>=')
					percent_pro = tmp1[0].strip('"')
					site_num1 = tmp1[1].split('sites')[0].strip()
					tmp2 = fs.readline().strip().split('=')
					xx_pro = tmp2[0].strip('"')
					site_num2 = tmp2[1].split('sites')[0].strip()
					mean_freq = fs.readline().strip().split('=')[1].strip('"').replace('%', '')
		
		for i, p in enumerate(paragraphs):
			if 'upRatio' in p.text:
				find_num = len(re.findall('upRatio', p.text))
				self.text_replace(p, ['upRatio'] * find_num, [str(self.fc)] * find_num)
			if 'downRatio' in p.text:
				find_num = len(re.findall('downRatio', p.text))
				self.text_replace(p, ['downRatio'] * find_num, [str(round(1 / self.fc, 2))] * find_num)
			if 'groupvs' in p.text:
				self.text_replace(p, ['groupvs'], [self.groupvs])
			if 'BP-TOP5' in p.text:
				if all([len(bp_top) == 0, len(mf_top) == 0, len(cc_top) == 0]):
					self.text_replace(p, ['，发生了显著性变化'], [''])
				if bp_top:
					self.text_replace(p, ['BP-TOP5'], [', '.join(bp_top)])
				else:
					self.text_replace(p, ['BP-TOP5等重要生物学过程'], ['无P值小于0.05的显著性生物学过程'])
			if 'MF-TOP5' in p.text:
				if mf_top:
					self.text_replace(p, ['MF-TOP5'], [', '.join(mf_top)])
				else:
					self.text_replace(p, ['MF-TOP5等分子功能'], ['无P值小于0.05的显著性分子功能'])
			if 'CC-TOP5' in p.text:
				if cc_top:
					self.text_replace(p, ['CC-TOP5'], [', '.join(cc_top)])
				else:
					self.text_replace(p, ['CC-TOP5等定位蛋白质'], ['无P值小于0.05的显著性定位蛋白'])
			if 'kegg-map2query-top5' in p.text:
				self.text_replace(p, ['kegg-map2query-top5'], [', '.join(map2query_top5)])
			if 'KeggEnrich-top5' in p.text:
				if self.keggEnrich_top5:
					self.text_replace(p, ['KeggEnrich-top5', 'KeggEnrich-top1'], [', '.join(self.keggEnrich_top5), self.pathway])
				else:
					if self.pathway:
						self.text_replace(p, ['KeggEnrich-top5等重要通路发生了显著变化', 'KeggEnrich-top1'], ['该比较组无P值小于0.05的显著性富集通路', self.pathway])
			if 'Percentage' in p.text:
				self.text_replace(p, ['Percentage', 'Median Score'], [self.percentage, self.medianScore])
			if 'groupNum' in p.text:
				self.text_replace(p, ['groupNum', 'Num', 'total'], [groupNum, Num, total_num])

			if '[Statistic]' in p.text:
				statistic_file = [i for i in os.listdir('Evaluation') if re.search('statistic.png', i , re.I)]
				if statistic_file:
					self.insert_png(p, '[Statistic]', os.path.join('Evaluation', statistic_file[0]))
			if '[venn1]' in p.text:
				is_replace = self.insert_png(p, '[venn1]', os.path.join('Evaluation', 'Venn', '组内', f'venn_{self.groupvs.split("_vs_")[0]}.png'))
				if is_replace == False:
					p.clear()
					self.paragraph_format(p.add_run('该项目无组内韦恩图'), size = 10.5, family = u'微软雅黑')
			if '[venn2]' in p.text:
					if venn2:
						self.insert_png(p, '[venn2]', os.path.join('Evaluation', 'Venn', '组间', f'{venn2[0]}'))
					else:
						p.clear()
						self.paragraph_format(p.add_run('该项目无法进行组间韦恩图绘制'), size = 10.5, family = u'微软雅黑')
						self.delete_paragraph(paragraphs, [i + 1])
			if '[pro_diff]' in p.text:
				self.insert_png(p, '[pro_diff]', os.path.join('Evaluation', 'dif_display.png'))
			if '[volcano]' in p.text:
				is_replace = self.insert_png(p, '[volcano]', os.path.join('Evaluation', 'VolcanoPlot', f'Volcano_Plot_{self.groupvs}.png'))
				if is_replace == False:
					self.text_replace(paragraphs[i - 1], ['，如下图'], [''])
					p.clear()
					self.paragraph_format(p.add_run('该项目结果无法进行火山图绘制。'), size = 10.5, family = u'微软雅黑')
					self.delete_paragraph(paragraphs, list(range(i + 1, i + 6)))
			if '[cluster]' in p.text:
				if os.path.exists(os.path.join(self.groupvs, 'CLUSTER')):
					self.insert_png(p, '[cluster]', os.path.join(self.groupvs, 'CLUSTER', 'cluster1.png'))
				elif os.path.exists(os.path.join(self.groupvs, 'cluster')):
					self.insert_png(p, '[cluster]', os.path.join(self.groupvs, 'cluster', 'cluster1.png'))
			if '[Subcellular_Localization]' in p.text:
				self.insert_png(p, '[Subcellular_Localization]', os.path.join(self.groupvs, 'cello', 'Subcellular_Localization.png'))
			if '[TopDomainStat]' in p.text:
				self.insert_png(p, '[TopDomainStat]', os.path.join(self.groupvs, 'domain', 'TopDomainStat.png'))
			if '[Domain_Enrichment]' in p.text:
				self.insert_png(p, '[Domain_Enrichment]', os.path.join(self.groupvs, 'domain', 'Domain_Enrichment.png'))
			if '[GOLevel2]' in p.text:
				self.insert_png(p, '[GOLevel2]', os.path.join(self.groupvs, 'go', 'GOLevel2.png'))
			if '[BP_Enrichment]' in p.text:
				self.insert_png(p, '[BP_Enrichment]', os.path.join(self.groupvs, 'go', 'BP_Enrichment.png'))
			if '[CC_Enrichment]' in p.text:
				self.insert_png(p, '[CC_Enrichment]', os.path.join(self.groupvs, 'go', 'CC_Enrichment.png'))
			if '[MF_Enrichment]' in p.text:
				self.insert_png(p, '[MF_Enrichment]', os.path.join(self.groupvs, 'go', 'MF_Enrichment.png'))
			if '[BP_DAG]' in p.text:
				self.insert_png(p, '[BP_DAG]', os.path.join(self.groupvs, 'go', 'BP_DAG.png'))
			if '[CC_DAG]' in p.text:
				self.insert_png(p, '[CC_DAG]', os.path.join(self.groupvs, 'go', 'CC_DAG.png'))
			if '[MF_DAG]' in p.text:
				self.insert_png(p, '[MF_DAG]', os.path.join(self.groupvs, 'go', 'MF_DAG.png'))
			if '[TopMapStat]' in p.text:
				self.insert_png(p, '[TopMapStat]', os.path.join(self.groupvs, 'kegg', 'TopMapStat.png'))
			if '[KEGG_Enrichment]' in p.text:
				self.insert_png(p, '[KEGG_Enrichment]', os.path.join(self.groupvs, 'kegg', 'KEGG_Enrichment.png'))
			if '[kegg_pathway]' in p.text:
				self.insert_png(p, '[kegg_pathway]', os.path.join(self.groupvs, 'kegg', 'map', f'{keggID}.png'))
			if '[ppi]' in p.text:
				self.insert_png(p, '[ppi]', os.path.join(self.groupvs, 'ppi', 'ppi.png'))
			if '[Module_ppi]' in p.text:
				is_replace = self.insert_png(p, '[Module_ppi]', os.path.join(self.groupvs, 'ppi', 'Module_ppi.png'))
				if is_replace == False:
					paragraphs[i - 1].clear()
					self.paragraph_format(paragraphs[i - 1].add_run('该比较组PPI结果无法进行蛋白网络模块分析。'), size = 10.5, family = u'微软雅黑')
					self.delete_paragraph(paragraphs, [i, i + 1])
			if '[mass_error]' in p.text:
				mass_error_file = [i for i in os.listdir('Evaluation') if re.search('mass.*.png', i, re.I)]
				if mass_error_file:
					self.insert_png(p, '[mass_error]', os.path.join('Evaluation', mass_error_file[0]))
				else:
					if '图3' in paragraphs[i -2].text:
						self.text_replace(paragraphs[i -2], ['如下图1所示，所有鉴定肽段的质量偏差主要分布在10ppm以内，说明鉴定结果准确可靠。然后', '图2', '图3'], ['', '图1', '图2'], family_ch=u'微软雅黑', family_en=u'微软雅黑')
						self.delete_paragraph(paragraphs, list(range(i - 1, i + 5)))
						self.text_replace(paragraphs[i + 5], ['5.2.2'], ['5.2.1'], size=12, bold=True)
						self.text_replace(paragraphs[i + 11], ['5.2.3'], ['5.2.2'], size=12, bold=True)
			if '[Andromeda_Score_Distribution]' in p.text:
				score_file = [i for i in os.listdir('Evaluation') if re.search('score.*.png', i, re.I)]
				if score_file:
					self.insert_png(p, '[Andromeda_Score_Distribution]', os.path.join('Evaluation', score_file[0]))
			if '[Ratio_Distribution]' in p.text:
				self.insert_png(p, '[Ratio_Distribution]', os.path.join('Evaluation', f'Ratio_Distribution_{self.groupvs}.png'))
			if '[MW_Distribution]' in p.text:
				self.insert_png(p, '[MW_Distribution]', os.path.join('Evaluation', 'Molecular_Weight_Distribution.png'))
			if '[pI_Distribution]' in p.text:
				self.insert_png(p, '[pI_Distribution]', os.path.join('Evaluation', 'pI_Distribution.png'))
			if '[PepLength_Distribution]' in p.text:
				pepLen_file = [i for i in os.listdir('Evaluation') if re.search('length', i, re.I)]
				if pepLen_file:
					self.insert_png(p, '[PepLength_Distribution]', os.path.join('Evaluation', pepLen_file[0]))
			if '[Protein_Sequence_Coverage_Distribution]' in p.text:
				self.insert_png(p, '[Protein_Sequence_Coverage_Distribution]', os.path.join('Evaluation', 'Protein_Sequence_Coverage_Distribution.png'))
			if '[Peptide_Count_Distribution]' in p.text:
				self.insert_png(p, '[Peptide_Count_Distribution]', os.path.join('Evaluation', 'Peptide_Count_Distribution.png'))

			### 磷酸化####
			if 'pro_pep_site' in p.text:
				pro_pep_site = '、'.join([self.statistic_list[4], self.statistic_list[2], self.statistic_list[0]])
				self.text_replace(p, ['pro_pep_site', 'pho_num1', 'pho_num2', 'pho_num3'], [pro_pep_site, self.statistic_list[-1], self.statistic_list[3], self.statistic_list[1]])
			if 'percent_pro' in p.text:
				self.text_replace(p, ['percent_pro', 'site_num1', 'xx_pro', 'site_num2'], [percent_pro, site_num1, xx_pro, site_num2])
			if 'mean_freq' in p.text:
				self.text_replace(p, ['mean_freq'], [mean_freq])
			if 'group1' in p.text:
				if self.groupvs != 'nan':
					self.text_replace(p, ['group1', 'group2'], [self.groupvs.split('_vs_')[0], self.groupvs.split('_vs_')[1]])
			if '[PhosphorylatedSitesInPro]' in p.text:
				siteinpro_file = [i for i in os.listdir('Evaluation') if re.search('sitesinpro.png', i, re.I)]
				if siteinpro_file:
					self.insert_png(p, '[PhosphorylatedSitesInPro]', os.path.join('Evaluation', siteinpro_file[0]))
			if '[PhosphorylatedFrequency]' in p.text:
				frequency_file = [i for i in os.listdir('Evaluation') if re.search('Frequency.png', i, re.I)]
				if frequency_file:
					self.insert_png(p, '[PhosphorylatedFrequency]', os.path.join('Evaluation', frequency_file[0]))
			if '[Phospho_STY_Distribution]' in p.text:
				self.insert_png(p, '[Phospho_STY_Distribution]', os.path.join('Evaluation', 'Phospho_STY_Distribution.png'))
			if '[inner1_venn]' in p.text:
				is_replace = self.insert_png(p, '[inner1_venn]', os.path.join('Evaluation', 'Venn', '组内', f'venn_Pro_{self.groupvs.split("_vs_")[0]}.png'))
				if is_replace == False:
					p.clear()
					self.paragraph_format(p.add_run('该项目无法进行韦恩图绘制'), size = 10.5, family = u'微软雅黑')
			if '[inner2_venn]' in p.text:
				is_replace = self.insert_png(p, '[inner2_venn]', os.path.join('Evaluation', 'Venn', '组内', f'venn_Pep_{self.groupvs.split("_vs_")[0]}.png'))
				if is_replace == False:
					p.clear()
					self.paragraph_format(p.add_run('该项目无法进行韦恩图绘制'), size = 10.5, family = u'微软雅黑')
			if '[all_pro_venn]' in p.text:
				if os.path.exists(os.path.join('Evaluation', 'Venn', '组间', 'ProteinVenn')):
					pro_venn_file = [i for i in os.listdir(os.path.join('Evaluation', 'Venn', '组间', 'ProteinVenn')) if re.search('png', i)]
					if pro_venn_file:
						self.insert_png(p, '[all_pro_venn]', os.path.join('Evaluation', 'Venn', '组间', 'ProteinVenn', pro_venn_file[0]))
			if '[all_pep_venn]' in p.text:
				if os.path.exists(os.path.join('Evaluation', 'Venn', '组间', 'PeptideVenn')):
					pep_venn_file = [i for i in os.listdir(os.path.join('Evaluation', 'Venn', '组间', 'PeptideVenn')) if re.search('png', i)]
					if pep_venn_file:
						self.insert_png(p, '[all_pep_venn]', os.path.join('Evaluation', 'Venn', '组间', 'PeptideVenn', pep_venn_file[0]))
			if os.path.exists('motif'):
				if '[motif_count]' in p.text:
					self.insert_png(p, '[motif_count]', os.path.join('motif', 'motif_count.png'))
				if '[motif_fold]' in p.text:
					self.insert_png(p, '[motif_fold]', os.path.join('motif', 'motif_fold.png'))

				if '[motif_fold_all]' in p.text:
					if os.path.exists(os.path.join('motif', 'motif_fold.txt')):
						self.delete_paragraph(paragraphs, [i])
						motif_all = pd.read_csv(os.path.join('motif', 'motif_fold.txt'), sep='\t', header=None).head(6)
						motif_all_list = motif_all.iloc[:, 1].tolist()
						if len(motif_all_list) < 4:
							motif_all_table = document.add_table(rows=2,cols=len(motif_all_list),style='Table Grid')
						elif len(motif_all_list) == 4:
							motif_all_table = document.add_table(rows=4,cols=2,style='Table Grid')
						else:
							motif_all_table = document.add_table(rows=4,cols=3,style='Table Grid')
						self.table_center(motif_all_table)
						self.move_table_after(motif_all_table, paragraphs[i - 1])
						for j in range(len(motif_all_list)):
							if j < 2:
								self.paragraph_format(motif_all_table.cell(0, j).paragraphs[0].add_run(f'Motif{str(j+1)}:{motif_all_list[j]}'), size = 10.5, family ='Arial',bold=True)
								run = motif_all_table.cell(1, j).paragraphs[0].add_run()
								run.add_picture(os.path.join('motif', f'{motif_all_list[j].replace(".", "x")}.png'), width=Inches(1.5))
							elif j == 2 or j == 3:
								if len(motif_all_list) < 4:
									self.paragraph_format(motif_all_table.cell(0, j).paragraphs[0].add_run(f'Motif{str(j+1)}:{motif_all_list[j]}'), size = 10.5, family ='Arial',bold=True)
									run = motif_all_table.cell(1, j).paragraphs[0].add_run()
									run.add_picture(os.path.join('motif', f'{motif_all_list[j].replace(".", "x")}.png'), width=Inches(1.5))
								elif len(motif_all_list) == 4:
									self.paragraph_format(motif_all_table.cell(2, j - 2).paragraphs[0].add_run(f'Motif{str(j+1)}:{motif_all_list[j]}'), size = 10.5, family ='Arial',bold=True)
									run = motif_all_table.cell(3, j - 2).paragraphs[0].add_run()
									run.add_picture(os.path.join('motif', f'{motif_all_list[j].replace(".", "x")}.png'), width=Inches(1.5))
								else:
									if j == 2:
										self.paragraph_format(motif_all_table.cell(0, j).paragraphs[0].add_run(f'Motif{str(j+1)}:{motif_all_list[j]}'), size = 10.5, family ='Arial',bold=True)
										run = motif_all_table.cell(1, j).paragraphs[0].add_run()
										run.add_picture(os.path.join('motif', f'{motif_all_list[j].replace(".", "x")}.png'), width=Inches(1.5))
									else:
										self.paragraph_format(motif_all_table.cell(2, j - 3).paragraphs[0].add_run(f'Motif{str(j+1)}:{motif_all_list[j]}'), size = 10.5, family ='Arial',bold=True)
										run = motif_all_table.cell(3, j - 3).paragraphs[0].add_run()
										run.add_picture(os.path.join('motif', f'{motif_all_list[j].replace(".", "x")}.png'), width=Inches(1.5))
							else:
								self.paragraph_format(motif_all_table.cell(2, j - 3).paragraphs[0].add_run(f'Motif{str(j+1)}:{motif_all_list[j]}'), size = 10.5, family ='Arial',bold=True)
								run = motif_all_table.cell(3, j - 3).paragraphs[0].add_run()
								run.add_picture(os.path.join('motif', f'{motif_all_list[j].replace(".", "x")}.png'), width=Inches(1.5))
					else:
						p.clear()
						if 'p' in self.types:
							self.paragraph_format(p.add_run('磷酸化肽段在数据库中无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'a' in self.types:
							self.paragraph_format(p.add_run('乙酰化肽段在数据库中无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'g' in self.types:
							self.paragraph_format(p.add_run('泛素化肽段在数据库中无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'n' in self.types:
							self.paragraph_format(p.add_run('糖基化肽段在数据库中无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 's' in self.types:
							self.paragraph_format(p.add_run('琥珀酰化肽段在数据库中无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'y' in self.types:
							self.paragraph_format(p.add_run('酪氨酸磷酸化肽段在数据库中无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'm' in self.types:
							self.paragraph_format(p.add_run('丙二酰化肽段在数据库中无motif'), size = 10.5, family = u'微软雅黑', bold=True)

						self.delete_paragraph(paragraphs, list(range(i + 6, i + 13)))

				if '[motif_up]' in p.text:
					if os.path.exists(os.path.join(self.groupvs, 'motif', 'up', 'motif_fold.txt')):
						self.delete_paragraph(paragraphs, [i])
						motif_up = pd.read_csv(os.path.join(self.groupvs, 'motif', 'up', 'motif_fold.txt'), sep='\t', header=None).head(3)
						motif_up_list = motif_up.iloc[:, 1].tolist()
						motif_up_table = document.add_table(rows=2,cols=len(motif_up_list),style='Table Grid')
						self.table_center(motif_up_table)
						self.move_table_after(motif_up_table, paragraphs[i - 1])
						for j in range(len(motif_up_list)):
							self.paragraph_format(motif_up_table.cell(0, j).paragraphs[0].add_run(f'Motif{str(j+1)}:{motif_up_list[j]}'), size = 10.5, family ='Arial',bold=True)
							run = motif_up_table.cell(1, j).paragraphs[0].add_run()
							run.add_picture(os.path.join(self.groupvs, 'motif', 'up', f'{motif_up_list[j].replace(".", "x")}.png'), width=Inches(1.5))
					else:
						p.clear()
						if 'p' in self.types:
							self.paragraph_format(p.add_run('上调磷酸化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'a' in self.types:
							self.paragraph_format(p.add_run('上调乙酰化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'g' in self.types:
							self.paragraph_format(p.add_run('上调泛素化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'n' in self.types:
							self.paragraph_format(p.add_run('上调糖基化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 's' in self.types:
							self.paragraph_format(p.add_run('上调琥珀酰化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'y' in self.types:
							self.paragraph_format(p.add_run('上调酪氨酸磷酸化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'm' in self.types:
							self.paragraph_format(p.add_run('上调丙二酰化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
							
				if '[motif_down]' in p.text:
					if os.path.exists(os.path.join(self.groupvs, 'motif', 'down', 'motif_fold.txt')):
						self.delete_paragraph(paragraphs, [i])
						motif_down = pd.read_csv(os.path.join(self.groupvs, 'motif', 'down', 'motif_fold.txt'), sep='\t', header=None).head(3)
						motif_down_list = motif_down.iloc[:, 1].tolist()
						motif_down_table = document.add_table(rows=2,cols=len(motif_down_list),style='Table Grid')
						self.table_center(motif_down_table)
						self.move_table_after(motif_down_table, paragraphs[i - 1])
						for j in range(len(motif_down_list)):
							self.paragraph_format(motif_down_table.cell(0, j).paragraphs[0].add_run(f'Motif{str(j+1)}:{motif_down_list[j]}'), size = 10.5, family ='Arial',bold=True)
							run = motif_down_table.cell(1, j).paragraphs[0].add_run()
							run.add_picture(os.path.join(self.groupvs, 'motif', 'down', f'{motif_down_list[j].replace(".", "x")}.png'), width=Inches(1.5))
					else:
						p.clear()
						if 'p' in self.types:
							self.paragraph_format(p.add_run('下调磷酸化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'a' in self.types:
							self.paragraph_format(p.add_run('下调乙酰化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'g' in self.types:
							self.paragraph_format(p.add_run('下调泛素化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'n' in self.types:
							self.paragraph_format(p.add_run('下调糖基化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 's' in self.types:
							self.paragraph_format(p.add_run('下调琥珀酰化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'y' in self.types:
							self.paragraph_format(p.add_run('下调酪氨酸磷酸化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)
						if 'm' in self.types:
							self.paragraph_format(p.add_run('下调丙二酰化肽段无motif'), size = 10.5, family = u'微软雅黑', bold=True)


	def save(self):
		pass

